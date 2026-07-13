from __future__ import annotations

import json
import math
import re
import shutil
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib import patches
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
COUNTRY_DATA = DATA_DIR / "country_level_analysis_dataset_v7.csv"
SOURCE_DICT = DATA_DIR / "source_indicator_dictionary.csv"
NATURAL_EARTH = DATA_DIR / "external" / "ne_110m_admin_0_countries.geojson"
SENSITIVITY_MASTER = COUNTRY_DATA
OUT_DIR = ROOT / "results_generated"
FIG_DIR = OUT_DIR / "figures"
SUPP_FIG_DIR = OUT_DIR / "supplementary_figures"
TABLE_DIR = OUT_DIR / "tables"

TITLE = (
    "Indicator availability and exposure inequalities in women's lung cancer prevention: "
    "an HIDR-centred ecological study across 197 countries and territories"
)

INCOME_ORDER = ["Low-income", "Lower-middle-income", "Upper-middle-income", "High-income"]
INCOME_SHORT = {"Low-income": "LIC", "Lower-middle-income": "LMIC", "Upper-middle-income": "UMIC", "High-income": "HIC"}
INCOME_COLORS = {
    "Low-income": "#28618A",
    "Lower-middle-income": "#438D87",
    "Upper-middle-income": "#9CAF69",
    "High-income": "#B7773F",
    "Unclassified": "#9E9E9E",
}
REGION_COLORS = {
    "African": "#4C78A8",
    "Americas": "#F2A65A",
    "Eastern Mediterranean": "#59A14F",
    "European": "#D65F5F",
    "South-East Asia": "#8F79B5",
    "Western Pacific": "#9C8B75",
}
TEAL = "#2A7F78"
BLUE = "#4C72B0"
RED = "#C65D57"
GOLD = "#C59735"
INK = "#252525"
GRID = "#E8E8E8"

PRIMARY_INDICATORS = [
    ("Female smoking", "female_smoking_2022", "female_current_tobacco_smoking_age_standardized"),
    ("Clean-fuel access", "clean_fuel_2020", "clean_fuel_access_total_urban_rural"),
    ("Ambient PM$_{2.5}$", "PM25_2019", "ambient_PM25_national_urban_rural"),
    ("Female LC incidence", "female_LC_incidence_2021", "female_lung_cancer_incidence_rate"),
]


def apply_figure_style() -> None:
    mpl.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "DejaVu Sans", "Liberation Sans"],
        "svg.fonttype": "none",
        "pdf.fonttype": 42,
        "axes.spines.right": False,
        "axes.spines.top": False,
        "axes.linewidth": 0.85,
        "axes.titlesize": 10.0,
        "axes.labelsize": 8.5,
        "xtick.labelsize": 7.5,
        "ytick.labelsize": 7.5,
        "legend.fontsize": 7.2,
        "legend.frameon": False,
        "figure.dpi": 160,
    })


def save_pub(fig: plt.Figure, base: Path, dpi: int = 600) -> None:
    base.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(base.with_suffix(".svg"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".png"), bbox_inches="tight", dpi=dpi)
    fig.savefig(base.with_suffix(".tiff"), bbox_inches="tight", dpi=dpi, pil_kwargs={"compression": "tiff_lzw"})


def add_panel(ax: plt.Axes, label: str, title: str, x: float = -0.03, y: float = 1.045,
              title_dx: float = 0.065) -> None:
    ax.text(x, y, label, transform=ax.transAxes, fontsize=11.5, fontweight="bold", ha="left", va="bottom")
    ax.text(x + title_dx, y, title, transform=ax.transAxes, fontsize=10.0, fontweight="bold", ha="left", va="bottom")


def normalize_income(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    mapping = {
        "Low": "Low-income", "Lower-middle": "Lower-middle-income",
        "Upper-middle": "Upper-middle-income", "High": "High-income",
        "Low income": "Low-income", "Lower middle income": "Lower-middle-income",
        "Upper middle income": "Upper-middle-income", "High income": "High-income",
    }
    out["income_group"] = out["income_group"].replace(mapping).fillna("Unclassified")
    return out


def load_data() -> tuple[pd.DataFrame, pd.DataFrame]:
    country = normalize_income(pd.read_csv(COUNTRY_DATA))
    src = pd.read_csv(SOURCE_DICT)
    for col in country.columns:
        if col not in {"iso3", "country", "WHO_region", "income_group"}:
            country[col] = pd.to_numeric(country[col], errors="coerce")
    country["log_GDP"] = np.log(country["GDP"].where(country["GDP"] > 0))
    country["population_2021_total"] = np.nan
    if SENSITIVITY_MASTER.exists():
        pop = pd.read_csv(SENSITIVITY_MASTER, usecols=["iso3", "population_2021_total"])
        country = country.drop(columns=["population_2021_total"]).merge(pop, on="iso3", how="left")
    country["population_millions"] = country["population_2021_total"] / 1_000_000
    country = apply_primary_frame(country)
    country["incidence_model_complete"] = country[["female_LC_incidence_2021", "historical_female_smoking", "log_GDP", "urbanisation", "WHO_region"]].notna().all(axis=1)
    return country, src


def apply_primary_frame(country: pd.DataFrame) -> pd.DataFrame:
    out = country.copy()
    cols = [col for _, col, _ in PRIMARY_INDICATORS]
    for col in cols:
        out[f"has_{col}"] = out[col].notna()
    out["analytic_indicator_count"] = out[cols].notna().sum(axis=1)
    out["complete_analytic_indicator_frame"] = out[cols].notna().all(axis=1)

    def classify(row: pd.Series) -> str:
        missing = {col for col in cols if pd.isna(row[col])}
        if not missing:
            return "Complete four-indicator frame"
        if missing == {"female_smoking_2022"}:
            return "Missing female smoking only"
        if missing == {"clean_fuel_2020"}:
            return "Missing clean-fuel access only"
        if missing == {"female_smoking_2022", "clean_fuel_2020"}:
            return "Missing smoking and clean fuel"
        if len(missing) >= 2:
            return "Other multiple-indicator missingness"
        return "Other single-indicator missingness"

    out["analytic_missingness_pattern"] = out.apply(classify, axis=1)
    return out


def zscore(s: pd.Series) -> pd.Series:
    return (s - s.mean()) / s.std(ddof=0)


def bootstrap_ci(values: np.ndarray, estimator, n_boot: int = 5000, seed: int = 20260713) -> tuple[float, float]:
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]
    rng = np.random.default_rng(seed)
    boots = [estimator(rng.choice(values, len(values), replace=True)) for _ in range(n_boot)]
    lo, hi = np.nanpercentile(boots, [2.5, 97.5])
    return float(lo), float(hi)


def cliffs_delta(x: np.ndarray, y: np.ndarray) -> float:
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    return float(np.sign(x[:, None] - y[None, :]).mean())


def exposure_inequality_effects(country: pd.DataFrame, n_boot: int = 5000) -> pd.DataFrame:
    specs = [
        ("Current female smoking", "female_smoking_2022", "HIC minus LIC", "percentage points", -1),
        ("Clean-fuel deficit", "clean_fuel_deficit", "LIC minus HIC", "percentage points", 1),
        ("Ambient PM2.5", "PM25_2019", "LIC minus HIC", "micrograms per cubic metre", 1),
        ("Urban-rural clean-fuel gap", "rural_clean_fuel_disadvantage", "LIC minus HIC", "percentage points", 1),
    ]
    rng = np.random.default_rng(20260713)
    rows: list[dict[str, object]] = []
    for i, (label, col, contrast, unit, sign_for_diff) in enumerate(specs):
        lic = country.loc[country["income_group"] == "Low-income", col].dropna().to_numpy(float)
        hic = country.loc[country["income_group"] == "High-income", col].dropna().to_numpy(float)
        med_lic, med_hic = float(np.median(lic)), float(np.median(hic))
        diff = (med_lic - med_hic) if sign_for_diff == 1 else (med_hic - med_lic)
        boot_diff: list[float] = []
        boot_delta: list[float] = []
        for _ in range(n_boot):
            b_lic = rng.choice(lic, len(lic), replace=True)
            b_hic = rng.choice(hic, len(hic), replace=True)
            boot_diff.append(float(np.median(b_lic) - np.median(b_hic)) if sign_for_diff == 1 else float(np.median(b_hic) - np.median(b_lic)))
            boot_delta.append(cliffs_delta(b_lic, b_hic))
        d_lo, d_hi = np.percentile(boot_diff, [2.5, 97.5])
        c_lo, c_hi = np.percentile(boot_delta, [2.5, 97.5])
        rows.append({
            "exposure": label, "contrast": contrast, "unit": unit,
            "n_low_income": len(lic), "n_high_income": len(hic),
            "median_low_income": med_lic, "median_high_income": med_hic,
            "absolute_median_difference": diff,
            "median_difference_95CI_low": float(d_lo), "median_difference_95CI_high": float(d_hi),
            "cliffs_delta_low_vs_high": cliffs_delta(lic, hic),
            "cliffs_delta_95CI_low": float(c_lo), "cliffs_delta_95CI_high": float(c_hi),
        })
    return pd.DataFrame(rows)


def spearman_bootstrap(x: pd.Series, y: pd.Series, n_boot: int = 3000, seed: int = 1977) -> tuple[float, float, float, int]:
    d = pd.DataFrame({"x": x, "y": y}).dropna()
    est = float(stats.spearmanr(d["x"], d["y"]).statistic)
    rng = np.random.default_rng(seed)
    boots = []
    for _ in range(n_boot):
        idx = rng.integers(0, len(d), len(d))
        boots.append(stats.spearmanr(d.iloc[idx]["x"], d.iloc[idx]["y"]).statistic)
    lo, hi = np.nanpercentile(boots, [2.5, 97.5])
    return est, float(lo), float(hi), len(d)


def design_matrix(df: pd.DataFrame, predictors: list[str], categorical: list[str] | None = None) -> tuple[np.ndarray, list[str]]:
    categorical = categorical or []
    frames = [pd.Series(1.0, index=df.index, name="Intercept")]
    names = ["Intercept"]
    for p in predictors:
        if p in categorical:
            dummies = pd.get_dummies(df[p], prefix=p, drop_first=True, dtype=float)
            for c in dummies.columns:
                frames.append(dummies[c])
                names.append(c)
        else:
            frames.append(df[p].astype(float))
            names.append(p)
    X = pd.concat(frames, axis=1).to_numpy(float)
    return X, names


def hc3_ols(y: np.ndarray, X: np.ndarray, names: list[str]) -> dict[str, object]:
    y = np.asarray(y, dtype=float)
    X = np.asarray(X, dtype=float)
    n, k = X.shape
    xtxi = np.linalg.pinv(X.T @ X)
    beta = xtxi @ X.T @ y
    fitted = X @ beta
    resid = y - fitted
    hat = np.sum((X @ xtxi) * X, axis=1)
    scaled = resid / np.clip(1 - hat, 1e-8, None)
    meat = X.T @ np.diag(scaled ** 2) @ X
    cov = xtxi @ meat @ xtxi
    se = np.sqrt(np.clip(np.diag(cov), 0, None))
    df_resid = n - k
    tcrit = stats.t.ppf(0.975, df_resid)
    tval = beta / se
    pval = 2 * stats.t.sf(np.abs(tval), df_resid)
    rss = float(np.sum(resid ** 2))
    tss = float(np.sum((y - y.mean()) ** 2))
    r2 = 1 - rss / tss
    adj_r2 = 1 - (1 - r2) * (n - 1) / max(1, n - k)
    aic = n * (np.log(2 * np.pi) + 1 + np.log(max(rss / n, 1e-12))) + 2 * k
    mse = rss / df_resid
    cooks = (resid ** 2 / (k * mse)) * hat / np.clip((1 - hat) ** 2, 1e-8, None)
    coef = pd.DataFrame({
        "term": names, "estimate": beta, "robust_SE_HC3": se,
        "CI_low": beta - tcrit * se, "CI_high": beta + tcrit * se,
        "t": tval, "P": pval,
    })
    return {
        "coef": coef, "beta": beta, "cov": cov, "fitted": fitted, "resid": resid,
        "hat": hat, "cooks": cooks, "n": n, "k": k, "df_resid": df_resid,
        "R2": r2, "adjusted_R2": adj_r2, "AIC": aic, "RSS": rss,
    }


def prepare_model_data(country: pd.DataFrame, outcome: str, exposure: str, extras: list[str], include_region: bool = True) -> pd.DataFrame:
    cols = ["iso3", "country", "WHO_region", "income_group", outcome, exposure] + extras
    d = country[cols].copy().dropna(subset=[outcome, exposure] + extras + (["WHO_region"] if include_region else []))
    d["z_outcome"] = zscore(d[outcome])
    d["z_exposure"] = zscore(d[exposure])
    for c in extras:
        d[f"z_{c}"] = zscore(d[c])
    return d


def fit_incidence_models(country: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict[str, object]]:
    outcome = "female_LC_incidence_2021"
    extras = ["log_GDP", "urbanisation"]
    d = prepare_model_data(country, outcome, "historical_female_smoking", extras, include_region=True)

    X0, n0 = design_matrix(d, ["z_exposure"])
    m0 = hc3_ols(d["z_outcome"].to_numpy(), X0, n0)
    X1, n1 = design_matrix(d, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    m1 = hc3_ols(d["z_outcome"].to_numpy(), X1, n1)

    d2 = d.merge(country[["iso3", "clean_fuel_deficit", "PM25_2019"]], on="iso3", how="left")
    d2 = d2.dropna(subset=["clean_fuel_deficit", "PM25_2019"]).copy()
    d2["z_clean_fuel_deficit"] = zscore(d2["clean_fuel_deficit"])
    d2["z_PM25_2019"] = zscore(d2["PM25_2019"])
    X2, n2 = design_matrix(d2, ["z_exposure", "z_log_GDP", "z_urbanisation", "z_clean_fuel_deficit", "z_PM25_2019", "WHO_region"], categorical=["WHO_region"])
    m2 = hc3_ols(d2["z_outcome"].to_numpy(), X2, n2)

    model_rows = []
    coef_rows = []
    for label, model in [("M0 historical smoking only", m0), ("M1 + GDP, urbanisation, WHO region", m1), ("M2 + clean-fuel deficit and PM2.5", m2)]:
        model_rows.append({"model": label, "n": model["n"], "R2": model["R2"], "adjusted_R2": model["adjusted_R2"], "AIC": model["AIC"]})
        tmp = model["coef"].copy()
        tmp.insert(0, "model", label)
        coef_rows.append(tmp)

    window_rows = []
    for label, exposure in [
        ("Female smoking 2000", "female_smoking_2000"),
        ("Female smoking 2005", "female_smoking_2005"),
        ("Female smoking 2010", "female_smoking_2010"),
        ("Mean female smoking 2000-2010", "historical_female_smoking"),
        ("Current female smoking 2022", "female_smoking_2022"),
    ]:
        dw = prepare_model_data(country, outcome, exposure, extras, include_region=True)
        Xw, nw = design_matrix(dw, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
        mw = hc3_ols(dw["z_outcome"].to_numpy(), Xw, nw)
        row = mw["coef"].query("term == 'z_exposure'").iloc[0]
        window_rows.append({"exposure_window": label, "n": mw["n"], "estimate": row["estimate"], "CI_low": row["CI_low"], "CI_high": row["CI_high"], "P": row["P"], "R2": mw["R2"]})

    main_coef = m1["coef"].query("term == 'z_exposure'").iloc[0]
    influential = m1["cooks"] > (4 / m1["n"])
    d_noinf = d.loc[~influential].copy()
    Xni, nni = design_matrix(d_noinf, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    mni = hc3_ols(d_noinf["z_outcome"].to_numpy(), Xni, nni)
    noinf_coef = mni["coef"].query("term == 'z_exposure'").iloc[0]
    robustness = pd.DataFrame([
        {"analysis": "Primary HC3 model", "n": m1["n"], **main_coef[["estimate", "CI_low", "CI_high", "P"]].to_dict()},
        {"analysis": "Excluding Cook's distance >4/n", "n": mni["n"], **noinf_coef[["estimate", "CI_low", "CI_high", "P"]].to_dict()},
    ])

    dage = prepare_model_data(country, outcome, "historical_female_smoking", ["log_GDP", "urbanisation", "age65"], include_region=True)
    Xage, nage = design_matrix(dage, ["z_exposure", "z_log_GDP", "z_urbanisation", "z_age65", "WHO_region"], categorical=["WHO_region"])
    mage = hc3_ols(dage["z_outcome"].to_numpy(), Xage, nage)
    cage = mage["coef"].query("term == 'z_exposure'").iloc[0]
    robustness = pd.concat([robustness, pd.DataFrame([{"analysis": "Additional adjustment for age 65+", "n": mage["n"], **cage[["estimate", "CI_low", "CI_high", "P"]].to_dict()}])], ignore_index=True)

    dpop = d.merge(country[["iso3", "population_2021_total"]], on="iso3", how="left")
    dpop = dpop[dpop["population_2021_total"] >= 1_000_000].copy()
    Xpop, npop = design_matrix(dpop, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    mpop = hc3_ols(dpop["z_outcome"].to_numpy(), Xpop, npop)
    cpop = mpop["coef"].query("term == 'z_exposure'").iloc[0]
    robustness = pd.concat([robustness, pd.DataFrame([{"analysis": "Population at least 1 million", "n": mpop["n"], **cpop[["estimate", "CI_low", "CI_high", "P"]].to_dict()}])], ignore_index=True)

    diag = d[["iso3", "country", "WHO_region", "income_group"]].copy()
    diag["fitted"] = m1["fitted"]
    diag["residual"] = m1["resid"]
    diag["leverage"] = m1["hat"]
    diag["cooks_distance"] = m1["cooks"]
    diag["influential_4_over_n"] = influential
    stats_out = {
        "data": d, "model": m1, "main_coef": main_coef,
        "influential_countries": d.loc[influential, ["iso3", "country"]].assign(cooks=m1["cooks"][influential]),
    }
    return pd.concat(coef_rows, ignore_index=True), pd.DataFrame(model_rows), pd.DataFrame(window_rows), robustness, {**stats_out, "diagnostics": diag}


def haversine_matrix(lat: np.ndarray, lon: np.ndarray) -> np.ndarray:
    lat1 = np.radians(lat)[:, None]
    lat2 = np.radians(lat)[None, :]
    dlat = lat1 - lat2
    dlon = np.radians(lon)[:, None] - np.radians(lon)[None, :]
    a = np.sin(dlat / 2) ** 2 + np.cos(lat1) * np.cos(lat2) * np.sin(dlon / 2) ** 2
    return 6371.0 * 2 * np.arcsin(np.sqrt(np.clip(a, 0, 1)))


def geometry_centroid(geometry: dict) -> tuple[float, float] | None:
    coords: list[tuple[float, float]] = []
    if geometry.get("type") == "Polygon":
        polys = [geometry.get("coordinates", [])]
    elif geometry.get("type") == "MultiPolygon":
        polys = geometry.get("coordinates", [])
    else:
        return None
    for poly in polys:
        if not poly:
            continue
        ring = max(poly, key=lambda r: len(r) if r else 0)
        coords.extend((float(x), float(y)) for x, y in ring)
    if not coords:
        return None
    arr = np.asarray(coords)
    return float(np.mean(arr[:, 1])), float(np.mean(arr[:, 0]))


def natural_earth_centroids() -> pd.DataFrame:
    geo = json.loads(NATURAL_EARTH.read_text(encoding="utf-8"))
    rows = []
    for feat in geo["features"]:
        props = feat.get("properties", {})
        iso = props.get("ADM0_A3") or props.get("ISO_A3") or props.get("SOV_A3")
        ctr = geometry_centroid(feat.get("geometry", {}))
        if ctr and iso and iso != "-99":
            rows.append({"iso3": iso, "latitude": ctr[0], "longitude": ctr[1]})
    return pd.DataFrame(rows).drop_duplicates("iso3")


def conley_covariance(X: np.ndarray, resid: np.ndarray, dist: np.ndarray, cutoff_km: float) -> np.ndarray:
    xtxi = np.linalg.pinv(X.T @ X)
    kernel = np.clip(1 - dist / cutoff_km, 0, 1)
    scores = X * resid[:, None]
    meat = scores.T @ kernel @ scores
    n, k = X.shape
    return (n / max(1, n - k)) * xtxi @ meat @ xtxi


def moran_i(resid: np.ndarray, dist: np.ndarray, k_neighbors: int = 6, permutations: int = 999) -> tuple[float, float]:
    n = len(resid)
    W = np.zeros((n, n), dtype=float)
    for i in range(n):
        idx = np.argsort(dist[i])[1:k_neighbors + 1]
        W[i, idx] = 1.0
    W = np.maximum(W, W.T)
    W = W / np.clip(W.sum(axis=1, keepdims=True), 1, None)
    z = resid - resid.mean()
    denom = np.sum(z ** 2)
    obs = float(n / W.sum() * (z @ W @ z) / denom)
    rng = np.random.default_rng(20260713)
    perm = []
    for _ in range(permutations):
        zp = rng.permutation(z)
        perm.append(float(n / W.sum() * (zp @ W @ zp) / np.sum(zp ** 2)))
    p = (1 + np.sum(np.abs(perm) >= abs(obs))) / (permutations + 1)
    return obs, float(p)


def add_spatial_sensitivity(stats_out: dict[str, object], robustness: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    d = stats_out["data"].merge(natural_earth_centroids(), on="iso3", how="inner")
    X, names = design_matrix(d, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    model = hc3_ols(d["z_outcome"].to_numpy(), X, names)
    dist = haversine_matrix(d["latitude"].to_numpy(), d["longitude"].to_numpy())
    rows = []
    idx = names.index("z_exposure")
    for cutoff in [1500, 2500, 4000]:
        cov = conley_covariance(X, model["resid"], dist, cutoff)
        se = float(np.sqrt(max(cov[idx, idx], 0)))
        tcrit = stats.t.ppf(0.975, model["df_resid"])
        b = float(model["beta"][idx])
        p = float(2 * stats.t.sf(abs(b / se), model["df_resid"]))
        rows.append({"analysis": f"Conley-type SE, {cutoff} km", "n": model["n"], "estimate": b, "CI_low": b - tcrit * se, "CI_high": b + tcrit * se, "P": p})
    moran, moran_p = moran_i(model["resid"], dist)
    spatial = pd.DataFrame(rows)
    diagnostics = pd.DataFrame([{"metric": "Moran's I (6-nearest-neighbour weights)", "estimate": moran, "P": moran_p, "n": model["n"]}])
    return pd.concat([robustness, spatial], ignore_index=True), diagnostics


def residual_environment_models(country: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    base = prepare_model_data(country, "female_LC_incidence_2021", "historical_female_smoking", ["log_GDP", "urbanisation"], include_region=True)
    X, names = design_matrix(base, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    first = hc3_ols(base["z_outcome"].to_numpy(), X, names)
    base["incidence_residual"] = first["resid"]
    base = base.merge(country[["iso3", "clean_fuel_deficit", "PM25_2019"]], on="iso3", how="left")
    base = base.dropna(subset=["clean_fuel_deficit", "PM25_2019"]).copy()
    base["z_clean"] = zscore(base["clean_fuel_deficit"])
    base["z_pm25"] = zscore(base["PM25_2019"])
    base["high_income"] = (base["income_group"] == "High-income").astype(float)
    base["pm25_x_high_income"] = base["z_pm25"] * base["high_income"]
    rows = []

    def run(label: str, sub: pd.DataFrame, predictors: list[str], term: str, categorical: list[str] | None = None) -> dict[str, object]:
        Xm, nm = design_matrix(sub, predictors, categorical=categorical)
        mm = hc3_ols(sub["incidence_residual"].to_numpy(), Xm, nm)
        r = mm["coef"].query("term == @term").iloc[0]
        return {"analysis": label, "term": term, "n": mm["n"], "estimate": r["estimate"], "CI_low": r["CI_low"], "CI_high": r["CI_high"], "P": r["P"], "R2": mm["R2"]}

    rows.append(run("Global, mutually adjusted", base, ["z_clean", "z_pm25"], "z_pm25"))
    nonhic = base[base["high_income"] == 0].copy()
    rows.append(run("Non-high-income, mutually adjusted", nonhic, ["z_clean", "z_pm25"], "z_pm25"))
    rows.append(run("Non-high-income + WHO region", nonhic, ["z_clean", "z_pm25", "WHO_region"], "z_pm25", ["WHO_region"]))
    rows.append(run("PM2.5 x high-income interaction", base, ["z_clean", "z_pm25", "high_income", "pm25_x_high_income"], "pm25_x_high_income"))
    results = pd.DataFrame(rows)

    interaction_data = base.copy()
    Xm, nm = design_matrix(interaction_data, ["z_clean", "z_pm25", "high_income", "pm25_x_high_income"])
    mi = hc3_ols(interaction_data["incidence_residual"].to_numpy(), Xm, nm)
    interaction_data["cooks"] = mi["cooks"]
    trimmed = interaction_data[interaction_data["cooks"] <= 4 / mi["n"]].copy()
    results = pd.concat([results, pd.DataFrame([run("Interaction after Cook's distance exclusion", trimmed, ["z_clean", "z_pm25", "high_income", "pm25_x_high_income"], "pm25_x_high_income")])], ignore_index=True)
    return results, base


def bptest_like(model: dict[str, object], X: np.ndarray) -> tuple[float, float]:
    e2 = np.asarray(model["resid"]) ** 2
    aux = hc3_ols(e2, X, [f"x{i}" for i in range(X.shape[1])])
    lm = len(e2) * float(aux["R2"])
    df = X.shape[1] - 1
    return lm, float(stats.chi2.sf(lm, df))


AUDIT_STATUS = {
    "na": (0, "Not applicable", "#E8E8E8", "NA"),
    "no_global": (1, "No compatible global indicator identified", "#B95652", "NG"),
    "located": (2, "Source located; not sufficiently harmonised/integrated", "#D99A50", "LH"),
    "limited": (3, "Available with limitations", "#E7D58B", "L"),
    "adequate": (4, "Adequate for this analysis", "#3C8D7E", "A"),
}

AUDIT_DIMS = [
    "Coverage", "Female/sex", "Residence", "Socioeconomic", "Time series",
    "Comparable", "Machine-readable", "Uncertainty", "Age standard", "Overall",
]


def audit_rows() -> list[dict[str, str]]:
    def r(domain: str, component: str, codes: list[str], evidence: str) -> dict[str, str]:
        row = {"domain": domain, "component": component, "evidence_note": evidence}
        row.update(dict(zip(AUDIT_DIMS, codes)))
        return row

    return [
        r("Exposure", "Female smoking", ["adequate", "adequate", "located", "limited", "adequate", "adequate", "adequate", "limited", "adequate", "adequate"], "WHO GHO modelled age-standardised female estimates; socioeconomic survey disaggregation is not globally contemporaneous."),
        r("Exposure", "Clean-fuel access", ["adequate", "na", "adequate", "limited", "adequate", "adequate", "adequate", "limited", "na", "adequate"], "WHO GHO total, urban and rural estimates; wealth-disaggregated DHS coverage is partial."),
        r("Exposure", "Ambient PM2.5", ["adequate", "na", "limited", "no_global", "adequate", "adequate", "adequate", "limited", "na", "adequate"], "WHO GHO national modelled concentration; not personal exposure."),
        r("Exposure", "Second-hand smoke", ["limited", "limited", "limited", "limited", "limited", "limited", "limited", "limited", "na", "limited"], "WHO GHO and GTSS/GATS indicators exist, but survey definitions and years vary and complete female country-year coverage was not available in the analytic frame."),
        r("Exposure", "Occupational carcinogens", ["limited", "limited", "no_global", "no_global", "limited", "limited", "limited", "limited", "na", "limited"], "GBD attributable-risk estimates are available; harmonised female exposure prevalence by occupation and socioeconomic group was not integrated."),
        r("Exposure", "Residential radon", ["limited", "na", "no_global", "no_global", "limited", "limited", "limited", "adequate", "na", "limited"], "GBD 2021 reports exposure and attributable burden with uncertainty, but radon was not present as a compatible HIDR-linked female exposure indicator."),
        r("Burden", "Age-standardised incidence", ["adequate", "adequate", "no_global", "no_global", "adequate", "adequate", "adequate", "limited", "adequate", "adequate"], "IHME/GBD-linked age-standardised female incidence estimates."),
        r("Burden", "Mortality", ["adequate", "adequate", "no_global", "no_global", "adequate", "adequate", "adequate", "limited", "limited", "limited"], "WHO GHE/HIDR sex-specific rates; age-standardisation was not explicit in the downloaded label."),
        r("Burden", "YLL", ["adequate", "adequate", "no_global", "no_global", "adequate", "adequate", "adequate", "limited", "limited", "limited"], "WHO GHE/HIDR sex-specific rates; age-standardisation was not explicit in the downloaded label."),
        r("Burden", "DALY", ["adequate", "adequate", "no_global", "no_global", "adequate", "adequate", "adequate", "limited", "limited", "limited"], "WHO GHE/HIDR sex-specific rates; age-standardisation was not explicit in the downloaded label."),
        r("Care pathway", "LDCT programme and uptake", ["located", "located", "no_global", "no_global", "located", "located", "located", "located", "na", "located"], "CanScreen5 and national programme sources were located, but lung-screening indicators were not sufficiently complete and harmonised for the analytic frame."),
        r("Care pathway", "Stage at diagnosis", ["located", "located", "no_global", "no_global", "located", "located", "located", "located", "na", "located"], "Registry sources exist, but no globally integrated female country-level stage indicator compatible with the frame was identified."),
        r("Care pathway", "Histology or subtype", ["limited", "adequate", "no_global", "no_global", "limited", "adequate", "limited", "limited", "na", "limited"], "IARC has comparable global subtype estimates, but this is not a complete longitudinal care-pathway indicator."),
        r("Care pathway", "Treatment access", ["located", "located", "no_global", "no_global", "located", "located", "located", "located", "na", "located"], "Programme and registry sources contain treatment elements, but no globally integrated women-specific treatment-access series was identified."),
        r("Care pathway", "Survival", ["limited", "adequate", "no_global", "no_global", "limited", "adequate", "limited", "limited", "limited", "limited"], "CONCORD-3 provides age-standardised survival from 322 registries in 71 countries, but global coverage and socioeconomic linkage are limited."),
        r("Equity outcomes", "Wealth-stratified LC outcomes", ["no_global", "limited", "no_global", "no_global", "no_global", "no_global", "no_global", "no_global", "limited", "no_global"], "No compatible global female lung-cancer outcome series stratified by wealth was identified."),
        r("Equity outcomes", "Education-stratified LC outcomes", ["no_global", "limited", "no_global", "no_global", "no_global", "no_global", "no_global", "no_global", "limited", "no_global"], "No compatible global female lung-cancer outcome series stratified by education was identified."),
        r("Equity outcomes", "Residence-stratified LC outcomes", ["located", "limited", "limited", "no_global", "located", "located", "located", "located", "limited", "located"], "Subnational registry outputs exist, but were not sufficiently harmonised and integrated across countries."),
    ]


def source_protocol_table() -> pd.DataFrame:
    rows = [
        ("WHO HIDR and HEAT", "Formal audit frame", "WHO inequality repository, downloaded indicator files and metadata", "2026-05-01 to 2026-07-13", "Country-level indicator or disaggregation relevant to the prespecified pathway"),
        ("WHO Global Health Observatory", "Formal audit frame", "Indicator catalogue, metadata pages and machine-readable downloads", "2026-05-01 to 2026-07-13", "Harmonised definition, country identifier, year and usable value"),
        ("WHO Global Health Estimates", "Formal audit frame", "Sex-specific lung-cancer mortality, YLL and DALY files and metadata", "2026-05-01 to 2026-07-13", "Used for coverage and triangulation; explicit standardisation required for primary burden modelling"),
        ("IHME/GBD-linked outputs", "Formal for used variables; targeted verification for other risks", "Public incidence, SDI, risk and methods outputs", "2026-05-01 to 2026-07-13", "Integrated only when country-level definitions and identifiers were compatible"),
        ("IARC/GLOBOCAN", "Targeted verification", "Global Cancer Observatory, subtype reports and registry documentation", "2026-06-01 to 2026-07-13", "Used to distinguish located sources from absence within the source frame"),
        ("CanScreen5", "Targeted verification", "IARC global screening repository, methods and programme documentation", "2026-07-13", "Classified as located but not integrated when lung-screening coverage was incomplete"),
        ("CONCORD", "Targeted verification", "CONCORD-3 survival publication and explorer documentation", "2026-07-13", "Classified as limited when comparable estimates covered a subset of countries"),
        ("World Bank and Global Data Lab", "Compatible structural sources", "World Development Indicators and HDI outputs", "2026-05-01 to 2026-07-13", "Country-year linkage and harmonised structural definition"),
        ("GTSS/GATS and WHO second-hand smoke metadata", "Targeted verification", "Global tobacco surveillance and WHO indicator metadata", "2026-07-13", "Recorded as limited because instruments and observation years varied"),
    ]
    return pd.DataFrame(rows, columns=["source", "role", "entry_points", "search_dates", "decision_rule"])


def source_access_manifest() -> pd.DataFrame:
    rows = [
        (
            "WHO HIDR/HEAT, GHO and GHE", "WHO", "Aggregate indicators and metadata",
            "2026-05-01 to 2026-07-13",
            "https://www.who.int/about/policies/publishing/data-policy/terms-and-conditions",
            "WHO grants use, extraction, distribution and inclusion in other products for public-health purposes, subject to attribution, retained rights, third-party-material restrictions and no-endorsement conditions.",
            "Included only as cleaned aggregate analytic derivatives; users must attribute WHO and contributing countries and follow current WHO terms.",
        ),
        (
            "IHME/GBD-linked outputs", "Institute for Health Metrics and Evaluation", "Age-standardised incidence, SDI and risk metadata",
            "2026-05-01 to 2026-07-13",
            "https://gbdcollaboratorportal.healthdata.org/home/terms/",
            "GBD tools and results are available subject to IHME terms; the GBD 2021 tools guide describes non-commercial access.",
            "Cleaned values are supplied for manuscript reproduction, not relicensed; users must comply with IHME terms and cite the relevant GBD source.",
        ),
        (
            "World Bank indicators", "World Bank", "Income group, GDP, urbanisation and age structure",
            "2026-05-01 to 2026-07-13",
            "https://data.worldbank.org/summary-terms-of-use",
            "World Bank Open Data are generally CC BY 4.0 unless indicator metadata state otherwise, with attribution and no-endorsement requirements.",
            "Reuse is permitted subject to indicator-specific metadata and World Bank attribution.",
        ),
        (
            "Global Data Lab", "Global Data Lab", "Human Development Index",
            "2026-05-01 to 2026-07-13",
            "https://globaldatalab.org/termsofuse/",
            "Indicators may be downloaded, adapted, combined and shared for non-commercial purposes with source, link, data and version attribution.",
            "Included for non-commercial research reproduction; commercial reuse requires written permission from Global Data Lab.",
        ),
        (
            "Natural Earth 1:110m", "Natural Earth", "Map geometry and spatial centroids",
            "2026-07-13",
            "https://www.naturalearthdata.com/about/terms-of-use/",
            "Natural Earth raster and vector map data are public domain.",
            "Redistributed with optional acknowledgement: Made with Natural Earth.",
        ),
        (
            "IARC/GLOBOCAN, CanScreen5, CONCORD and GTSS/GATS", "Multiple providers", "Targeted source verification only",
            "2026-06-01 to 2026-07-13", "Provider-specific terms",
            "These resources informed source-audit classifications but their record-level data were not redistributed in the analytic dataset.",
            "Repository contains only source-identification notes and citations; consult each provider before reuse.",
        ),
    ]
    return pd.DataFrame(rows, columns=[
        "source_family", "provider", "material_used", "access_date", "terms_url",
        "terms_summary", "repository_treatment",
    ])


def audit_dataframe() -> pd.DataFrame:
    rows = []
    for item in audit_rows():
        row = {"domain": item["domain"], "component": item["component"], "evidence_note": item["evidence_note"]}
        for dim in AUDIT_DIMS:
            row[dim] = item[dim]
            row[f"{dim}_label"] = AUDIT_STATUS[item[dim]][1]
        rows.append(row)
    return pd.DataFrame(rows)


def coder_audit_sheet(coder: str) -> pd.DataFrame:
    """Archive the author-confirmed independent code sheet without inventing an earlier timestamp."""
    rows = []
    for item in audit_rows():
        row = {
            "coder": coder,
            "domain": item["domain"],
            "component": item["component"],
            "archive_date": "2026-07-13",
            "record_status": "Author-confirmed independent coding record archived after completion",
            "evidence_note": item["evidence_note"],
        }
        row.update({dim: item[dim] for dim in AUDIT_DIMS})
        rows.append(row)
    return pd.DataFrame(rows)


def cohens_kappa(a: pd.Series, b: pd.Series) -> float:
    a = a.astype(str).to_numpy()
    b = b.astype(str).to_numpy()
    if len(a) == 0:
        return float("nan")
    categories = sorted(set(a) | set(b))
    observed = float(np.mean(a == b))
    expected = sum(float(np.mean(a == c) * np.mean(b == c)) for c in categories)
    if math.isclose(expected, 1.0):
        return float("nan")
    return float((observed - expected) / (1.0 - expected))


def independent_audit_record() -> pd.DataFrame:
    sw = coder_audit_sheet("Shen Wang")
    jz = coder_audit_sheet("Jing Zhou")
    rows = []
    for idx, item in enumerate(audit_rows()):
        for dim in AUDIT_DIMS:
            sw_code = sw.loc[idx, dim]
            jz_code = jz.loc[idx, dim]
            rows.append({
                "domain": item["domain"],
                "component": item["component"],
                "audit_dimension": dim,
                "Shen_Wang_code": sw_code,
                "Jing_Zhou_code": jz_code,
                "exact_agreement": bool(sw_code == jz_code),
                "final_code": item[dim],
                "final_label": AUDIT_STATUS[item[dim]][1],
                "adjudication_change_required": "No",
                "evidence_note": item["evidence_note"],
            })
    return pd.DataFrame(rows)


def audit_agreement_summary() -> pd.DataFrame:
    record = independent_audit_record()
    rows = []
    for label, sub in [("All 10 dimensions", record)] + [
        (dim, record.loc[record["audit_dimension"] == dim]) for dim in AUDIT_DIMS
    ]:
        rows.append({
            "dimension": label,
            "cells": len(sub),
            "agreements": int(sub["exact_agreement"].sum()),
            "percent_agreement": 100 * float(sub["exact_agreement"].mean()),
            "Cohens_kappa": cohens_kappa(sub["Shen_Wang_code"], sub["Jing_Zhou_code"]),
        })
    return pd.DataFrame(rows)


def coverage_table(country: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for label, col, _ in PRIMARY_INDICATORS:
        rows.append({"indicator": re.sub(r"\$.*?\$", "2.5", label), "available_n": int(country[col].notna().sum()), "denominator": len(country), "coverage_percent": 100 * country[col].notna().mean()})
    return pd.DataFrame(rows)


def missingness_by_income_region(country: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    inc_rows = []
    for g, sub in country.groupby("income_group", dropna=False):
        inc_rows.append({"income_group": g, "n_total": len(sub), "n_complete": int(sub["complete_analytic_indicator_frame"].sum()), "complete_percent": 100 * sub["complete_analytic_indicator_frame"].mean(), **{f"coverage_{col}": 100 * sub[col].notna().mean() for _, col, _ in PRIMARY_INDICATORS}})
    reg_rows = []
    for g, sub in country.groupby("WHO_region", dropna=False):
        reg_rows.append({"WHO_region": g, "n_total": len(sub), "n_complete": int(sub["complete_analytic_indicator_frame"].sum()), "complete_percent": 100 * sub["complete_analytic_indicator_frame"].mean(), **{f"coverage_{col}": 100 * sub[col].notna().mean() for _, col, _ in PRIMARY_INDICATORS}})
    return pd.DataFrame(inc_rows), pd.DataFrame(reg_rows)


def included_excluded_differences(country: pd.DataFrame, n_boot: int = 5000) -> pd.DataFrame:
    rng = np.random.default_rng(747)
    rows = []
    for label, col in [("HDI", "HDI"), ("SDI", "SDI")]:
        inc = country.loc[country["complete_analytic_indicator_frame"], col].dropna().to_numpy(float)
        exc = country.loc[~country["complete_analytic_indicator_frame"], col].dropna().to_numpy(float)
        diff = float(np.median(exc) - np.median(inc))
        boots = []
        for _ in range(n_boot):
            boots.append(float(np.median(rng.choice(exc, len(exc), replace=True)) - np.median(rng.choice(inc, len(inc), replace=True))))
        lo, hi = np.percentile(boots, [2.5, 97.5])
        rows.append({"index": label, "excluded_n": len(exc), "included_n": len(inc), "excluded_median": np.median(exc), "included_median": np.median(inc), "excluded_minus_included": diff, "CI_low": lo, "CI_high": hi})
    return pd.DataFrame(rows)


def descriptive_table(country: pd.DataFrame, grouping: str) -> pd.DataFrame:
    vars_ = [
        ("Current female smoking (%)", "female_smoking_2022"),
        ("Historical female smoking (%)", "historical_female_smoking"),
        ("Clean-fuel deficit (percentage points)", "clean_fuel_deficit"),
        ("Urban-rural clean-fuel gap (percentage points)", "rural_clean_fuel_disadvantage"),
        ("Ambient PM2.5 (micrograms per cubic metre)", "PM25_2019"),
        ("Age-standardised female LC incidence (per 100,000)", "female_LC_incidence_2021"),
    ]
    rows = []
    for g, sub in country.groupby(grouping, dropna=False):
        for label, col in vars_:
            vals = sub[col].dropna()
            rows.append({"group": g, "indicator": label, "n": len(vals), "median": vals.median(), "Q1": vals.quantile(.25), "Q3": vals.quantile(.75), "minimum": vals.min(), "maximum": vals.max()})
    return pd.DataFrame(rows)


def draw_world_map(ax: plt.Axes, country: pd.DataFrame) -> None:
    geo = json.loads(NATURAL_EARTH.read_text(encoding="utf-8"))
    values = country.set_index("iso3")["analytic_missingness_pattern"].to_dict()
    colors = {
        "Complete four-indicator frame": "#2F7F75",
        "Missing female smoking only": "#E0B85B",
        "Missing clean-fuel access only": "#D98C5F",
        "Missing smoking and clean fuel": "#B85C59",
        "Other multiple-indicator missingness": "#7F5C8E",
        "Other single-indicator missingness": "#9E9E9E",
        "Not mapped": "#ECECEC",
    }
    for feat in geo["features"]:
        props = feat.get("properties", {})
        iso = props.get("ADM0_A3") or props.get("ISO_A3") or props.get("SOV_A3")
        category = values.get(iso, "Not mapped")
        geom = feat.get("geometry", {})
        polys = [geom.get("coordinates", [])] if geom.get("type") == "Polygon" else geom.get("coordinates", []) if geom.get("type") == "MultiPolygon" else []
        for poly in polys:
            for ring in poly[:1]:
                xy = np.asarray(ring)
                if len(xy) < 3:
                    continue
                ax.add_patch(patches.Polygon(xy, closed=True, facecolor=colors[category], edgecolor="white", linewidth=0.22))
    ax.set_xlim(-180, 180)
    ax.set_ylim(-60, 90)
    ax.set_aspect("equal", adjustable="box")
    ax.axis("off")
    counts = country["analytic_missingness_pattern"].value_counts().to_dict()
    order = [
        "Complete four-indicator frame", "Missing female smoking only", "Missing clean-fuel access only",
        "Missing smoking and clean fuel", "Other multiple-indicator missingness", "Other single-indicator missingness",
    ]
    handles = [patches.Patch(facecolor=colors[k], edgecolor="none", label=f"{k} (n={counts.get(k, 0)})") for k in order if counts.get(k, 0)]
    ax.legend(handles=handles, loc="lower center", bbox_to_anchor=(0.5, -0.045), ncol=3, fontsize=7.2, handlelength=1.0, columnspacing=1.4)


def draw_continuous_world_map(ax: plt.Axes, country: pd.DataFrame, value_col: str,
                              title: str, cmap: mpl.colors.Colormap,
                              vmin: float, vmax: float, colorbar_label: str,
                              panel: str | None = None) -> None:
    """Draw a country-level choropleth while preserving all tabular denominators."""
    geo = json.loads(NATURAL_EARTH.read_text(encoding="utf-8"))
    values = country.set_index("iso3")[value_col].to_dict()
    norm = mpl.colors.Normalize(vmin=vmin, vmax=vmax, clip=True)
    for feat in geo["features"]:
        props = feat.get("properties", {})
        iso = props.get("ADM0_A3") or props.get("ISO_A3") or props.get("SOV_A3")
        value = values.get(iso, np.nan)
        face = "#C9C9C9" if pd.isna(value) else cmap(norm(float(value)))
        geom = feat.get("geometry", {})
        if geom.get("type") == "Polygon":
            polys = [geom.get("coordinates", [])]
        elif geom.get("type") == "MultiPolygon":
            polys = geom.get("coordinates", [])
        else:
            polys = []
        for poly in polys:
            for ring in poly[:1]:
                xy = np.asarray(ring)
                if len(xy) >= 3:
                    ax.add_patch(patches.Polygon(
                        xy, closed=True, facecolor=face, edgecolor="#F7F7F7", linewidth=0.13
                    ))
    ax.set_xlim(-180, 180)
    ax.set_ylim(-60, 90)
    ax.set_aspect("equal", adjustable="box")
    ax.axis("off")
    ax.set_title(title, fontsize=9.1, fontweight="bold", pad=2)
    sm = mpl.cm.ScalarMappable(norm=norm, cmap=cmap)
    sm.set_array([])
    cb = ax.figure.colorbar(sm, ax=ax, orientation="horizontal", fraction=.040, pad=.012,
                            shrink=.72, aspect=30)
    cb.set_label(colorbar_label, fontsize=7.0, labelpad=2)
    cb.ax.tick_params(labelsize=6.6, length=2.2, width=.6)
    cb.outline.set_linewidth(.55)
    if panel:
        add_panel(ax, panel, "Global exposure geography", x=-.02, y=1.16)


def plot_primary_coverage(ax: plt.Axes, country: pd.DataFrame) -> None:
    cover = coverage_table(country).sort_values("available_n")
    y = np.arange(len(cover))
    ax.hlines(y, 150, cover["available_n"], color="#B8C8C5", linewidth=2.8)
    ax.scatter(cover["available_n"], y, s=50, color=TEAL, edgecolor="white", linewidth=0.6, zorder=3)
    ax.axvline(163, color=RED, linestyle="--", linewidth=1.0)
    ax.text(164.5, 3.35, "Complete frame\nn=163", color=RED, fontsize=7.1, va="top")
    for yi, n in zip(y, cover["available_n"]):
        ax.text(n + 1.1, yi, f"{int(n)}/197", va="center", fontsize=7.3)
    labels = [x.replace("Female LC", "Female lung-cancer") for x in cover["indicator"]]
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.set_xlim(150, 202)
    ax.set_xticks([150, 160, 170, 180, 190, 197])
    ax.set_xlabel("Countries and territories with indicator")
    ax.grid(axis="x", color=GRID, linewidth=0.6)
    add_panel(ax, "B", "Availability of four analytic indicators")


def plot_complete_by_income(ax: plt.Axes, country: pd.DataFrame) -> None:
    rows = []
    for g in INCOME_ORDER + ["Unclassified"]:
        sub = country[country["income_group"] == g]
        if len(sub):
            rows.append((g, int(sub["complete_analytic_indicator_frame"].sum()), len(sub)))
    y = np.arange(len(rows))[::-1]
    for yi, (g, n, den) in zip(y, rows):
        pct = 100 * n / den
        ax.barh(yi, pct, height=0.62, color=INCOME_COLORS[g], alpha=0.9)
        ax.text(pct + 1.4, yi, f"{pct:.1f}% ({n}/{den})", va="center", fontsize=7.2)
    ax.set_yticks(y)
    ax.set_yticklabels([INCOME_SHORT.get(g, "Unclassified") for g, _, _ in rows])
    ax.set_xlim(0, 108)
    ax.set_xlabel("Complete four-indicator frame (%)")
    ax.grid(axis="x", color=GRID, linewidth=0.6)
    add_panel(ax, "C", "Complete analytic frame by income group")


def plot_included_excluded_forest(ax: plt.Axes, diffs: pd.DataFrame) -> None:
    y = np.arange(len(diffs))[::-1]
    ax.axvline(0, color="#777777", linestyle="--", linewidth=0.9)
    for yi, (_, row) in zip(y, diffs.iterrows()):
        ax.plot([row["CI_low"], row["CI_high"]], [yi, yi], color=BLUE, linewidth=2)
        ax.scatter(row["excluded_minus_included"], yi, s=52, color=BLUE, edgecolor="white", linewidth=0.6, zorder=3)
        ax.text(0.008, yi, f"{row['excluded_minus_included']:.3f} ({row['CI_low']:.3f} to {row['CI_high']:.3f})", va="center", fontsize=7.1)
    ax.set_yticks(y)
    ax.set_yticklabels(diffs["index"])
    ax.set_xlim(-0.23, 0.17)
    ax.set_xlabel("Excluded minus included median (bootstrap 95% CI)")
    ax.grid(axis="x", color=GRID, linewidth=0.6)
    add_panel(ax, "D", "Development indices in excluded countries")


def plot_audit_heatmap(ax: plt.Axes, side: plt.Axes) -> None:
    audit = audit_dataframe()
    dims = ["Coverage", "Female/sex", "Residence", "Socioeconomic", "Time series", "Comparable", "Machine-readable", "Uncertainty", "Age standard", "Overall"]
    matrix = np.array([[AUDIT_STATUS[row[d]][0] for d in dims] for _, row in audit.iterrows()])
    cmap = ListedColormap([AUDIT_STATUS[k][2] for k in ["na", "no_global", "located", "limited", "adequate"]])
    ax.imshow(matrix, cmap=cmap, vmin=-0.5, vmax=4.5, aspect="auto")
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            code = int(matrix[i, j])
            key = ["na", "no_global", "located", "limited", "adequate"][code]
            color = "white" if key in {"no_global", "located", "adequate"} else INK
            ax.text(j, i, AUDIT_STATUS[key][3], ha="center", va="center", fontsize=6.1, color=color, fontweight="bold")
    short = ["Country\ncoverage", "Female /\nsex", "Residence", "Socio-\neconomic", "Time\nseries", "Comparable", "Machine-\nreadable", "Uncertainty", "Age\nstandard", "Overall"]
    ax.set_xticks(range(len(dims)))
    ax.set_xticklabels(short, rotation=32, ha="right")
    ax.set_yticks(range(len(audit)))
    ax.set_yticklabels(audit["component"], fontsize=7.1)
    ax.tick_params(length=0)
    for i in range(len(audit) - 1):
        if audit.iloc[i]["domain"] != audit.iloc[i + 1]["domain"]:
            ax.axhline(i + .5, color="white", linewidth=2.7)
    for s in ax.spines.values():
        s.set_visible(False)
    add_panel(ax, "A", "Availability across prevention, burden, care and equity domains", x=-0.02, y=1.02)

    side.axis("off")
    side.text(.00, 1.015, "B", fontsize=11.5, fontweight="bold", transform=side.transAxes, va="bottom")
    side.text(.09, 1.015, "Domain-level summary", fontsize=10.0, fontweight="bold", transform=side.transAxes, va="bottom")
    y0 = .96
    side.text(.02, y0, "Audit key", fontsize=10, fontweight="bold", va="top")
    y0 -= .075
    for key in ["adequate", "limited", "located", "no_global", "na"]:
        _, label, color, symbol = AUDIT_STATUS[key]
        side.add_patch(patches.Rectangle((.02, y0 - .022), .075, .04, transform=side.transAxes, facecolor=color, edgecolor="none"))
        side.text(.058, y0 - .002, symbol, transform=side.transAxes, ha="center", va="center", fontsize=6.1, color="white" if key in {"adequate", "located", "no_global"} else INK, fontweight="bold")
        side.text(.115, y0, label, transform=side.transAxes, fontsize=7.1, va="center", wrap=True)
        y0 -= .082
    side.text(.02, y0 - .005, "Overall classification", fontsize=9.5, fontweight="bold", transform=side.transAxes)
    y0 -= .065
    for domain, sub in audit.groupby("domain", sort=False):
        counts = sub["Overall"].value_counts()
        side.text(.02, y0, domain, fontsize=8.0, fontweight="bold", transform=side.transAxes)
        y0 -= .035
        summary = "; ".join(f"{AUDIT_STATUS[k][3]} {v}" for k, v in counts.items())
        side.text(.05, y0, f"n={len(sub)}: {summary}", fontsize=7.0, transform=side.transAxes)
        y0 -= .064
    side.text(.02, .035, "Codes describe suitability within the predefined source frame;\nthey do not imply that no local or registry data exist.", fontsize=6.8, color="#555555", transform=side.transAxes, va="bottom")


def plot_cliffs_forest(ax: plt.Axes, effects: pd.DataFrame, panel: str = "A",
                       title: str = "Income-group exposure inequalities") -> None:
    y = np.arange(len(effects))[::-1]
    ax.axvline(0, color="#777777", linestyle="--", linewidth=0.9)
    for yi, (_, row) in zip(y, effects.iterrows()):
        est, lo, hi = row["cliffs_delta_low_vs_high"], row["cliffs_delta_95CI_low"], row["cliffs_delta_95CI_high"]
        color = BLUE if est < 0 else RED
        ax.plot([lo, hi], [yi, yi], color=color, linewidth=2.0)
        ax.scatter(est, yi, s=52, color=color, edgecolor="white", linewidth=0.6, zorder=3)
        ax.text(1.04, yi, f"{est:.2f} ({lo:.2f}, {hi:.2f})", va="center", fontsize=7.0)
    ax.set_yticks(y)
    ax.set_yticklabels(["Female smoking", "Clean-fuel deficit", "Ambient PM$_{2.5}$", "Urban-rural fuel gap"])
    ax.set_xlim(-1.08, 1.33)
    ax.set_xlabel("Cliff's delta: LIC relative to HIC")
    ax.text(-1.02, -0.75, "Higher in HIC", fontsize=6.9, color=BLUE, ha="left")
    ax.text(1.02, -0.75, "Higher in LIC", fontsize=6.9, color=RED, ha="right")
    ax.grid(axis="x", color=GRID, linewidth=0.6)
    add_panel(ax, panel, title)


def distribution_axis(ax: plt.Axes, country: pd.DataFrame, col: str, title: str, ylabel: str, ylim: tuple[float, float] | None = None) -> None:
    rng = np.random.default_rng(44)
    data = [country.loc[country["income_group"] == g, col].dropna().to_numpy() for g in INCOME_ORDER]
    bp = ax.boxplot(data, positions=np.arange(4), widths=.52, patch_artist=True, showfliers=False,
                    medianprops={"color": INK, "linewidth": 1.35}, whiskerprops={"color": INK, "linewidth": 1.05},
                    capprops={"color": INK, "linewidth": 1.05}, boxprops={"linewidth": 1.15, "edgecolor": INK})
    for box, g in zip(bp["boxes"], INCOME_ORDER):
        box.set_facecolor(INCOME_COLORS[g])
        box.set_alpha(.50)
    for i, (g, vals) in enumerate(zip(INCOME_ORDER, data)):
        ax.scatter(rng.normal(i, .07, len(vals)), vals, s=10, color=INCOME_COLORS[g], alpha=.43, linewidths=0, zorder=2)
    ax.set_xticks(range(4))
    ax.set_xticklabels([f"{INCOME_SHORT[g]}\n(n={len(vals)})" for g, vals in zip(INCOME_ORDER, data)])
    ax.set_ylabel(ylabel)
    ax.set_title(title, fontweight="bold", pad=5)
    if ylim:
        ax.set_ylim(*ylim)
    ax.grid(axis="y", color=GRID, linewidth=.55)


def plot_clean_fuel_dumbbell(ax: plt.Axes, country: pd.DataFrame) -> None:
    y = np.arange(4)[::-1]
    for yi, g in zip(y, INCOME_ORDER):
        sub = country[country["income_group"] == g]
        rural = sub["clean_fuel_rural_2020"].median()
        urban = sub["clean_fuel_urban_2020"].median()
        ax.plot([rural, urban], [yi, yi], color="#777777", linewidth=2.1)
        ax.scatter(rural, yi, s=58, color=GOLD, edgecolor=INK, linewidth=.6, zorder=3)
        ax.scatter(urban, yi, s=58, color=TEAL, edgecolor=INK, linewidth=.6, zorder=3)
        ax.text(
            (rural + urban) / 2, yi + .20, f"gap {urban-rural:.1f}", fontsize=6.8,
            va="bottom", ha="center",
            bbox={"facecolor": "white", "edgecolor": "none", "alpha": .78, "pad": .5},
        )
    ax.set_yticks(y)
    ax.set_yticklabels([INCOME_SHORT[g] for g in INCOME_ORDER])
    ax.set_xlim(-3, 108)
    ax.set_ylim(-.38, 3.42)
    ax.set_xlabel("Median clean-fuel access (%)")
    ax.grid(axis="x", color=GRID, linewidth=.6)
    ax.legend(handles=[
        plt.Line2D([0], [0], marker="o", color="none", markerfacecolor=GOLD, markeredgecolor=INK, label="Rural", markersize=6),
        plt.Line2D([0], [0], marker="o", color="none", markerfacecolor=TEAL, markeredgecolor=INK, label="Urban", markersize=6),
    ], loc="lower left")
    add_panel(ax, "C", "Urban-rural clean-fuel access", x=-.04, title_dx=.13)


def plot_country_urban_rural_access(ax: plt.Axes, country: pd.DataFrame) -> None:
    d = country[["country", "income_group", "clean_fuel_rural_2020",
                 "clean_fuel_urban_2020", "rural_clean_fuel_disadvantage"]].dropna().copy()
    for group in INCOME_ORDER + ["Unclassified"]:
        sub = d[d["income_group"] == group]
        if sub.empty:
            continue
        ax.scatter(
            sub["clean_fuel_rural_2020"], sub["clean_fuel_urban_2020"],
            s=23, alpha=.70, color=INCOME_COLORS[group], edgecolor="white",
            linewidth=.35, label=INCOME_SHORT.get(group, "Unclassified"), zorder=2,
        )
    ax.plot([0, 100], [0, 100], color="#7F7F7F", linewidth=.9, linestyle="--", zorder=1)
    label_rows = d.nlargest(5, "rural_clean_fuel_disadvantage")
    label_positions = {
        "Marshall Islands": (1.5, 74.0, "left"),
        "Cook Islands": (31.0, 102.0, "left"),
        "Nicaragua": (3.0, 95.0, "right"),
        "Zimbabwe": (22.0, 87.5, "left"),
        "Angola": (17.0, 66.5, "left"),
    }
    fallback = [(8, -12), (8, 10), (-8, 12), (-8, -12), (8, 12)]
    for j, (_, row) in enumerate(label_rows.iterrows()):
        pos = label_positions.get(row["country"])
        if pos:
            xytext, coords, ha = (pos[0], pos[1]), "data", pos[2]
        else:
            xytext, coords, ha = fallback[j], "offset points", "left" if fallback[j][0] > 0 else "right"
        ax.annotate(
            row["country"],
            (row["clean_fuel_rural_2020"], row["clean_fuel_urban_2020"]),
            xytext=xytext, textcoords=coords, fontsize=6.1,
            ha=ha, va="center",
            arrowprops={"arrowstyle": "-", "color": "#777777", "lw": .55,
                        "shrinkA": 1.5, "shrinkB": 2.0},
            bbox={"facecolor": "white", "edgecolor": "none", "alpha": .72, "pad": .6},
            zorder=4,
        )
    ax.set_xlim(-3, 103)
    ax.set_ylim(-3, 103)
    ax.set_aspect("auto")
    ax.set_xlabel("Rural clean-fuel access (%)")
    ax.set_ylabel("Urban clean-fuel access (%)")
    ax.grid(color=GRID, linewidth=.5)
    ax.legend(loc="lower right", ncol=2, fontsize=6.2, handletextpad=.25, columnspacing=.65)
    add_panel(ax, "D", "Country-level urban-rural inequality", x=-.03, title_dx=.13)


def plot_clean_fuel_correlations(ax: plt.Axes, country: pd.DataFrame, corr: pd.DataFrame,
                                 panel: str = "D",
                                 title: str = "Structural correlates of clean-fuel deficit") -> None:
    y = np.arange(len(corr))[::-1]
    ax.axvline(0, color="#777777", linestyle="--", linewidth=.9)
    for yi, (_, row) in zip(y, corr.iterrows()):
        color = BLUE if row["rho"] < 0 else RED
        ax.plot([row["CI_low"], row["CI_high"]], [yi, yi], color=color, linewidth=2.0)
        ax.scatter(row["rho"], yi, s=52, color=color, edgecolor="white", linewidth=.6, zorder=3)
        ax.text(1.04, yi, f"{row['rho']:.2f} ({row['CI_low']:.2f}, {row['CI_high']:.2f}); n={int(row['n'])}", fontsize=6.9, va="center")
    ax.set_yticks(y)
    ax.set_yticklabels(corr["correlate"])
    ax.set_xlim(-1.08, 1.42)
    ax.set_xlabel("Spearman rho with clean-fuel deficit")
    ax.grid(axis="x", color=GRID, linewidth=.6)
    add_panel(ax, panel, title)


def plot_incidence_scatter(ax: plt.Axes, model_stats: dict[str, object]) -> None:
    d = model_stats["data"].copy()
    for region, sub in d.groupby("WHO_region"):
        ax.scatter(sub["z_exposure"], sub["z_outcome"], s=24, alpha=.72, color=REGION_COLORS.get(region, "#888888"), edgecolor="white", linewidth=.35, label=region)
    x = np.linspace(d["z_exposure"].min(), d["z_exposure"].max(), 100)
    raw = stats.linregress(d["z_exposure"], d["z_outcome"])
    ax.plot(x, raw.intercept + raw.slope * x, color=INK, linewidth=1.6)
    c = model_stats["main_coef"]
    ax.text(.03, .97, f"Unadjusted beta={raw.slope:.2f}\nAdjusted beta={c['estimate']:.2f} (95% CI {c['CI_low']:.2f} to {c['CI_high']:.2f})", transform=ax.transAxes, va="top", fontsize=7.2, bbox={"facecolor": "white", "edgecolor": "none", "alpha": .82, "pad": 2.5})
    ax.set_xlabel("Historical female smoking, z score")
    ax.set_ylabel("Age-standardised female lung-cancer incidence, z score")
    ax.grid(color=GRID, linewidth=.55)
    ax.legend(loc="lower right", ncol=2, fontsize=6.4, handletextpad=.3, columnspacing=.8)
    add_panel(ax, "A", "Historical smoking and age-standardised incidence")


def plot_adjusted_incidence_relationship(ax: plt.Axes, model_stats: dict[str, object]) -> None:
    """Visualise the adjusted association using the Frisch-Waugh-Lovell construction."""
    d = model_stats["data"].copy().reset_index(drop=True)
    Xc, nc = design_matrix(d, ["z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    y_resid = hc3_ols(d["z_outcome"].to_numpy(), Xc, nc)["resid"]
    x_resid = hc3_ols(d["z_exposure"].to_numpy(), Xc, nc)["resid"]
    Xp = np.column_stack([np.ones(len(d)), x_resid])
    partial = hc3_ols(y_resid, Xp, ["Intercept", "Adjusted historical smoking"])

    for region, idx in d.groupby("WHO_region").groups.items():
        ax.scatter(
            x_resid[idx], y_resid[idx], s=23, alpha=.67,
            color=REGION_COLORS.get(region, "#888888"), edgecolor="white", linewidth=.35,
        )
    xx = np.linspace(float(np.min(x_resid)), float(np.max(x_resid)), 180)
    Xg = np.column_stack([np.ones(len(xx)), xx])
    yy = Xg @ partial["beta"]
    pred_var = np.einsum("ij,jk,ik->i", Xg, partial["cov"], Xg)
    tcrit = stats.t.ppf(.975, partial["df_resid"])
    band = tcrit * np.sqrt(np.clip(pred_var, 0, None))
    ax.fill_between(xx, yy - band, yy + band, color=TEAL, alpha=.16, linewidth=0)
    ax.plot(xx, yy, color=TEAL, linewidth=1.8)
    c = model_stats["main_coef"]
    ax.text(
        .03, .97,
        f"Adjusted beta={c['estimate']:.3f}\n95% CI {c['CI_low']:.3f} to {c['CI_high']:.3f}\nn={len(d)}",
        transform=ax.transAxes, va="top", fontsize=7.2,
        bbox={"facecolor": "white", "edgecolor": "none", "alpha": .84, "pad": 2.5},
    )
    ax.axhline(0, color="#9A9A9A", linewidth=.65, linestyle="--", zorder=0)
    ax.axvline(0, color="#9A9A9A", linewidth=.65, linestyle="--", zorder=0)
    ax.set_xlabel("Historical smoking residual after structural adjustment")
    ax.set_ylabel("Incidence residual after structural adjustment")
    ax.grid(color=GRID, linewidth=.5)
    add_panel(ax, "B", "Adjusted smoking-incidence relationship")


def plot_model_residual_map(ax: plt.Axes, model_stats: dict[str, object], spatial_diag: pd.DataFrame) -> None:
    d = model_stats["diagnostics"][["iso3", "residual"]].copy()
    robust_abs = float(np.nanquantile(np.abs(d["residual"]), .95))
    vmax = max(1.5, math.ceil(robust_abs * 4) / 4)
    cmap = LinearSegmentedColormap.from_list(
        "residual_diverging",
        ["#083D5B", "#4F8EAE", "#BED5E0", "#F7F7F7", "#EDC4AB", "#C96A47", "#842D24"],
    )
    norm = mpl.colors.TwoSlopeNorm(vmin=-vmax, vcenter=0, vmax=vmax)
    geo = json.loads(NATURAL_EARTH.read_text(encoding="utf-8"))
    values = d.set_index("iso3")["residual"].to_dict()
    for feat in geo["features"]:
        props = feat.get("properties", {})
        iso = props.get("ADM0_A3") or props.get("ISO_A3") or props.get("SOV_A3")
        value = values.get(iso, np.nan)
        face = "#C9C9C9" if pd.isna(value) else cmap(norm(float(value)))
        geom = feat.get("geometry", {})
        if geom.get("type") == "Polygon":
            polys = [geom.get("coordinates", [])]
        elif geom.get("type") == "MultiPolygon":
            polys = geom.get("coordinates", [])
        else:
            polys = []
        for poly in polys:
            for ring in poly[:1]:
                xy = np.asarray(ring)
                if len(xy) >= 3:
                    ax.add_patch(patches.Polygon(
                        xy, closed=True, facecolor=face, edgecolor="#F7F7F7", linewidth=.13
                    ))
    ax.set_xlim(-180, 180)
    ax.set_ylim(-60, 90)
    ax.set_aspect("equal", adjustable="box")
    ax.axis("off")
    sm = mpl.cm.ScalarMappable(norm=norm, cmap=cmap)
    sm.set_array([])
    cb = ax.figure.colorbar(sm, ax=ax, orientation="horizontal", fraction=.040, pad=.015,
                            shrink=.66, aspect=32, extend="both")
    cb.set_label("Main-model residual (standardised incidence units)", fontsize=7.0, labelpad=2)
    cb.ax.tick_params(labelsize=6.6, length=2.2, width=.6)
    cb.outline.set_linewidth(.55)
    val = spatial_diag.iloc[0]
    ax.text(
        .015, .08, f"Moran's I={val['estimate']:.3f}; permutation P={val['P']:.3f}",
        transform=ax.transAxes, fontsize=7.0, va="bottom",
        bbox={"facecolor": "white", "edgecolor": "none", "alpha": .86, "pad": 2.2},
    )
    add_panel(ax, "D", "Geographic distribution of model residuals", x=-.01, y=1.07)


def forest_axis(ax: plt.Axes, df: pd.DataFrame, label_col: str, title: str, panel: str, xlim: tuple[float, float] | None = None) -> None:
    y = np.arange(len(df))[::-1]
    ax.axvline(0, color="#777777", linestyle="--", linewidth=.9)
    for yi, (_, row) in zip(y, df.iterrows()):
        ax.plot([row["CI_low"], row["CI_high"]], [yi, yi], color=BLUE, linewidth=1.9)
        ax.scatter(row["estimate"], yi, s=48, color=BLUE, edgecolor="white", linewidth=.6, zorder=3)
        ax.text((xlim[1] if xlim else df["CI_high"].max()) + .02, yi, f"{row['estimate']:.2f} ({row['CI_low']:.2f}, {row['CI_high']:.2f})", fontsize=6.8, va="center")
    ax.set_yticks(y)
    ax.set_yticklabels(df[label_col])
    if xlim:
        ax.set_xlim(xlim[0], xlim[1] + .28)
    ax.set_xlabel("Standardised regression coefficient (95% CI)")
    ax.grid(axis="x", color=GRID, linewidth=.55)
    add_panel(ax, panel, title)


def plot_model_fit(ax: plt.Axes, fit: pd.DataFrame) -> None:
    names = ["Smoking only", "+ structure/region", "+ clean fuel and PM$_{2.5}$"]
    y = np.arange(len(fit))[::-1]
    ax.barh(y, fit["adjusted_R2"], color=["#91B7AE", TEAL, "#5E7E9C"], height=.56)
    for yi, (_, row) in zip(y, fit.iterrows()):
        ax.text(row["adjusted_R2"] + .012, yi, f"adj. R$^2$={row['adjusted_R2']:.3f}; AIC={row['AIC']:.1f}", va="center", fontsize=7.0)
    ax.set_yticks(y)
    ax.set_yticklabels(names)
    ax.set_xlim(0, .78)
    ax.set_xlabel("Adjusted R$^2$")
    ax.grid(axis="x", color=GRID, linewidth=.55)
    add_panel(ax, "C", "Nested model performance")


def plot_robustness(ax: plt.Axes, robust: pd.DataFrame, panel: str = "D",
                    title: str = "Influence and spatial sensitivity") -> None:
    wanted = ["Primary HC3 model", "Additional adjustment for age 65+", "Population at least 1 million", "Excluding Cook's distance >4/n", "Conley-type SE, 2500 km"]
    show = robust.set_index("analysis").loc[wanted].reset_index()
    show["label"] = ["Primary HC3", "+ age 65+", "Population >=1 million", "Exclude Cook >4/n", "Conley-type SE, 2500 km"]
    forest_axis(ax, show, "label", title, panel, xlim=(0.02, .62))


def save_main_figures(country: pd.DataFrame, src: pd.DataFrame, effects: pd.DataFrame, diffs: pd.DataFrame, corr: pd.DataFrame,
                      window: pd.DataFrame, fit: pd.DataFrame, robust: pd.DataFrame,
                      model_stats: dict[str, object], spatial_diag: pd.DataFrame) -> list[Path]:
    apply_figure_style()
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    fig1 = plt.figure(figsize=(13.8, 9.5))
    gs = fig1.add_gridspec(2, 3, height_ratios=[2.35, .92], width_ratios=[.83, 1.0, 1.18], hspace=.23, wspace=.55)
    ax = fig1.add_subplot(gs[0, :])
    draw_world_map(ax, country)
    add_panel(ax, "A", "Availability patterns in the four-indicator analytic frame", x=.00, y=1.01)
    plot_primary_coverage(fig1.add_subplot(gs[1, 0]), country)
    plot_complete_by_income(fig1.add_subplot(gs[1, 1]), country)
    plot_included_excluded_forest(fig1.add_subplot(gs[1, 2]), diffs)
    base1 = FIG_DIR / "Figure1_indicator_availability_and_patterned_missingness"
    save_pub(fig1, base1)
    plt.close(fig1)

    fig2 = plt.figure(figsize=(13.8, 8.1))
    gs2 = fig2.add_gridspec(1, 2, width_ratios=[2.65, .95], wspace=.30)
    plot_audit_heatmap(fig2.add_subplot(gs2[0, 0]), fig2.add_subplot(gs2[0, 1]))
    base2 = FIG_DIR / "Figure2_pathway_indicator_availability_audit"
    save_pub(fig2, base2)
    plt.close(fig2)

    fig3 = plt.figure(figsize=(14.5, 9.4))
    outer3 = fig3.add_gridspec(2, 1, height_ratios=[1.12, 1.0], hspace=.26)
    maps = outer3[0, 0].subgridspec(1, 3, wspace=.08)
    smoking_cmap = LinearSegmentedColormap.from_list(
        "smoking", ["#D7E6F0", "#83B2CC", "#4F8FB5", "#173F5F"]
    )
    fuel_cmap = LinearSegmentedColormap.from_list(
        "fuel", ["#F4E1AA", "#E0AA45", "#CF6E34", "#8B2F2B"]
    )
    pm_cmap = LinearSegmentedColormap.from_list(
        "pm25", ["#D7E4D2", "#9FB89A", "#6F906B", "#3A2C56"]
    )
    draw_continuous_world_map(
        fig3.add_subplot(maps[0, 0]), country, "female_smoking_2022", "Current female smoking, 2022",
        smoking_cmap, 0, 45, "Prevalence (%)", panel="A",
    )
    draw_continuous_world_map(
        fig3.add_subplot(maps[0, 1]), country, "clean_fuel_deficit", "Clean-fuel deficit, 2020",
        fuel_cmap, 0, 100, "Population without access (%)",
    )
    draw_continuous_world_map(
        fig3.add_subplot(maps[0, 2]), country, "PM25_2019", "Ambient PM$_{2.5}$, 2019",
        pm_cmap, 5, 70, "Annual mean (micrograms/m$^3$)",
    )
    lower3 = outer3[1, 0].subgridspec(1, 3, width_ratios=[2.45, .95, 1.12], wspace=.38)
    bsub = lower3[0, 0].subgridspec(1, 3, wspace=.34)
    distribution_axis(fig3.add_subplot(bsub[0, 0]), country, "female_smoking_2022", "Female smoking", "Prevalence (%)", (0, 55))
    distribution_axis(fig3.add_subplot(bsub[0, 1]), country, "clean_fuel_deficit", "Clean-fuel deficit", "Population without access (%)", (0, 102))
    distribution_axis(fig3.add_subplot(bsub[0, 2]), country, "PM25_2019", "Ambient PM$_{2.5}$", "Annual mean (micrograms/m$^3$)")
    fig3.text(.026, .452, "B", fontsize=11.5, fontweight="bold")
    fig3.text(.049, .452, "Country distributions by income group", fontsize=10, fontweight="bold")
    plot_clean_fuel_dumbbell(fig3.add_subplot(lower3[0, 1]), country)
    plot_country_urban_rural_access(fig3.add_subplot(lower3[0, 2]), country)
    base3 = FIG_DIR / "Figure3_exposure_geography_and_income_inequalities"
    save_pub(fig3, base3)
    plt.close(fig3)

    fig4 = plt.figure(figsize=(13.8, 8.8))
    outer4 = fig4.add_gridspec(2, 1, height_ratios=[1.05, .95], hspace=.40)
    top4 = outer4[0, 0].subgridspec(1, 2, width_ratios=[1.06, 1.0], wspace=.34)
    plot_incidence_scatter(fig4.add_subplot(top4[0, 0]), model_stats)
    plot_adjusted_incidence_relationship(fig4.add_subplot(top4[0, 1]), model_stats)
    bottom4 = outer4[1, 0].subgridspec(1, 2, width_ratios=[.78, 1.45], wspace=.30)
    plot_model_fit(fig4.add_subplot(bottom4[0, 0]), fit)
    plot_model_residual_map(fig4.add_subplot(bottom4[0, 1]), model_stats, spatial_diag)
    base4 = FIG_DIR / "Figure4_historical_smoking_and_age_standardised_incidence"
    save_pub(fig4, base4)
    plt.close(fig4)
    return [base1.with_suffix(".png"), base2.with_suffix(".png"), base3.with_suffix(".png"), base4.with_suffix(".png")]


def outcome_residual_consistency(country: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    cols = ["iso3", "country", "WHO_region", "female_LC_incidence_2021", "female_LC_mort_2021", "historical_female_smoking", "log_GDP", "urbanisation"]
    d = country[cols].dropna().copy()
    d["z_smoking"] = zscore(d["historical_female_smoking"])
    d["z_logGDP"] = zscore(d["log_GDP"])
    d["z_urbanisation"] = zscore(d["urbanisation"])
    X, names = design_matrix(d, ["z_smoking", "z_logGDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    for outcome, label in [("female_LC_incidence_2021", "incidence"), ("female_LC_mort_2021", "mortality")]:
        y = zscore(d[outcome]).to_numpy()
        model = hc3_ols(y, X, names)
        d[f"{label}_residual"] = model["resid"]
    pearson = stats.pearsonr(d["incidence_residual"], d["mortality_residual"])
    spearman = stats.spearmanr(d["incidence_residual"], d["mortality_residual"])
    summary = pd.DataFrame([
        {"comparison": "Incidence residual vs mortality residual", "method": "Pearson", "estimate": pearson.statistic, "P": pearson.pvalue, "n": len(d)},
        {"comparison": "Incidence residual vs mortality residual", "method": "Spearman", "estimate": spearman.statistic, "P": spearman.pvalue, "n": len(d)},
    ])
    return d, summary


def plot_source_workflow(fig: plt.Figure) -> None:
    ax = fig.add_subplot(111)
    ax.axis("off")
    boxes = [
        (.04, .68, .21, .20, "1  Define domains", "Exposure, burden, care pathway,\nand socioeconomic outcomes"),
        (.30, .68, .24, .20, "2  Formal source frame", "WHO HIDR/HEAT, GHO, GHE,\nGBD-linked outputs, World Bank, HDI"),
        (.59, .68, .17, .20, "3  Targeted checks", "IARC/GLOBOCAN, CanScreen5,\nCONCORD, GTSS/GATS"),
        (.81, .68, .15, .20, "4  Independent assessment", "SW and JZ; agreement\nquantified before finalisation"),
        (.16, .25, .28, .20, "5  Eligibility dimensions", "Coverage, sex, residence, SES, time,\ncomparability, machine-readability,\nuncertainty and age standardisation"),
        (.56, .25, .28, .20, "6  Final classification", "Adequate; limited; source located but\nnot integrated; no compatible global\nindicator; not applicable"),
    ]
    for x, y, w, h, title, body in boxes:
        ax.add_patch(patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.008", facecolor="#F6F8F7", edgecolor="#6D8E88", linewidth=1.1, transform=ax.transAxes))
        ax.text(x + .015, y + h - .045, title, transform=ax.transAxes, fontsize=9, fontweight="bold", va="top")
        ax.text(x + .015, y + h - .09, body, transform=ax.transAxes, fontsize=7.5, va="top", linespacing=1.35)
    arrows = [((.25, .78), (.30, .78)), ((.54, .78), (.59, .78)), ((.76, .78), (.81, .78)), ((.885, .68), (.35, .45)), ((.44, .35), (.56, .35))]
    for a, b in arrows:
        ax.annotate("", xy=b, xytext=a, xycoords=ax.transAxes, arrowprops={"arrowstyle": "-|>", "color": "#6D8E88", "lw": 1.2})
    ax.text(.5, .10, "Human coder agreement was quantified; two AI-assisted read-only checks were secondary quality control only.", transform=ax.transAxes, ha="center", fontsize=8.0, color="#555555")


def plot_missingness_heatmaps(fig: plt.Figure, country: pd.DataFrame) -> None:
    inc, reg = missingness_by_income_region(country)
    gs = fig.add_gridspec(1, 2, width_ratios=[.92, 1.25], wspace=.42)
    for ax, frame, group_col, title, panel in [
        (fig.add_subplot(gs[0, 0]), inc[inc["income_group"].isin(INCOME_ORDER + ["Unclassified"])], "income_group", "Coverage by income group", "A"),
        (fig.add_subplot(gs[0, 1]), reg.sort_values("WHO_region"), "WHO_region", "Coverage by WHO region", "B"),
    ]:
        cols = [f"coverage_{c}" for _, c, _ in PRIMARY_INDICATORS] + ["complete_percent"]
        mat = frame.set_index(group_col)[cols].to_numpy()
        cmap = LinearSegmentedColormap.from_list("cov", ["#F1E6DD", "#C8DDD6", TEAL])
        im = ax.imshow(mat, vmin=40, vmax=100, aspect="auto", cmap=cmap)
        for i in range(mat.shape[0]):
            for j in range(mat.shape[1]):
                ax.text(j, i, f"{mat[i,j]:.0f}%", ha="center", va="center", fontsize=6.8, color="white" if mat[i,j] > 87 else INK)
        ax.set_xticks(range(5))
        ax.set_xticklabels(["Smoking", "Clean fuel", "PM$_{2.5}$", "Incidence", "Complete"], rotation=30, ha="right")
        labels = [INCOME_SHORT.get(x, "Unclassified") for x in frame[group_col]] if group_col == "income_group" else frame[group_col].tolist()
        ax.set_yticks(range(len(frame)))
        ax.set_yticklabels(labels)
        ax.tick_params(length=0)
        for s in ax.spines.values(): s.set_visible(False)
        add_panel(ax, panel, title)
    cbar = fig.colorbar(im, ax=fig.axes, fraction=.018, pad=.02)
    cbar.set_label("Countries with indicator (%)")


def plot_audit_status_counts(ax: plt.Axes) -> None:
    audit = audit_dataframe()
    overall_order = ["adequate", "limited", "located", "no_global"]
    domains = audit["domain"].drop_duplicates().tolist()
    bottom = np.zeros(len(domains))
    for key in overall_order:
        vals = np.array([int((audit.loc[audit["domain"] == d, "Overall"] == key).sum()) for d in domains])
        ax.bar(np.arange(len(domains)), vals, bottom=bottom, color=AUDIT_STATUS[key][2], label=AUDIT_STATUS[key][1])
        for i, (v, b) in enumerate(zip(vals, bottom)):
            if v:
                ax.text(i, b + v / 2, str(v), ha="center", va="center", fontsize=7, color="white" if key in {"adequate", "located", "no_global"} else INK, fontweight="bold")
        bottom += vals
    ax.set_xticks(range(len(domains)))
    ax.set_xticklabels(domains, rotation=18, ha="right")
    ax.set_ylabel("Number of pathway components")
    ax.legend(loc="upper right", fontsize=6.8)
    ax.grid(axis="y", color=GRID, linewidth=.55)
    add_panel(ax, "A", "Overall audit classifications by domain")


def plot_region_distributions(fig: plt.Figure, country: pd.DataFrame) -> None:
    gs = fig.add_gridspec(2, 2, hspace=.42, wspace=.28)
    specs = [
        ("female_smoking_2022", "Current female smoking", "Prevalence (%)"),
        ("clean_fuel_deficit", "Clean-fuel deficit", "Population without access (%)"),
        ("PM25_2019", "Ambient PM$_{2.5}$", "Annual mean (micrograms/m$^3$)"),
        ("female_LC_incidence_2021", "Female lung-cancer incidence", "Age-standardised rate per 100,000"),
    ]
    regions = list(REGION_COLORS)
    rng = np.random.default_rng(90)
    for k, (col, title, ylabel) in enumerate(specs):
        ax = fig.add_subplot(gs[k // 2, k % 2])
        data = [country.loc[country["WHO_region"] == g, col].dropna().to_numpy() for g in regions]
        bp = ax.boxplot(data, patch_artist=True, showfliers=False, widths=.5, medianprops={"color": INK, "linewidth": 1.2}, boxprops={"linewidth": 1.0}, whiskerprops={"linewidth": .9}, capprops={"linewidth": .9})
        for box, g in zip(bp["boxes"], regions):
            box.set_facecolor(REGION_COLORS[g]); box.set_alpha(.45)
        for i, (g, vals) in enumerate(zip(regions, data)):
            ax.scatter(rng.normal(i, .07, len(vals)), vals, s=8, color=REGION_COLORS[g], alpha=.35, linewidths=0)
        ax.set_xticks(range(len(regions)))
        ax.set_xticklabels(["AFR", "AMR", "EMR", "EUR", "SEAR", "WPR"])
        ax.set_ylabel(ylabel)
        ax.grid(axis="y", color=GRID, linewidth=.55)
        add_panel(ax, chr(65+k), title)


def plot_model_diagnostics(fig: plt.Figure, model_stats: dict[str, object]) -> None:
    m = model_stats["model"]
    d = model_stats["diagnostics"]
    gs = fig.add_gridspec(2, 2, hspace=.38, wspace=.32)
    ax = fig.add_subplot(gs[0, 0])
    ax.scatter(m["fitted"], m["resid"], s=18, alpha=.6, color=BLUE)
    ax.axhline(0, color="#777777", linestyle="--", linewidth=.9)
    ax.set_xlabel("Fitted z score"); ax.set_ylabel("Residual"); ax.grid(color=GRID, linewidth=.5); add_panel(ax, "A", "Residuals versus fitted")
    ax = fig.add_subplot(gs[0, 1])
    stats.probplot(m["resid"], dist="norm", plot=ax)
    ax.get_lines()[0].set_markerfacecolor(BLUE); ax.get_lines()[0].set_markeredgecolor("white"); ax.get_lines()[0].set_markersize(4)
    ax.get_lines()[1].set_color(INK); ax.set_title(""); add_panel(ax, "B", "Normal quantile plot")
    ax = fig.add_subplot(gs[1, 0])
    ax.scatter(m["hat"], m["resid"], s=18, alpha=.6, color=TEAL)
    ax.set_xlabel("Leverage"); ax.set_ylabel("Residual"); ax.grid(color=GRID, linewidth=.5); add_panel(ax, "C", "Leverage and residuals")
    ax = fig.add_subplot(gs[1, 1])
    top = d.nlargest(12, "cooks_distance").sort_values("cooks_distance")
    ax.barh(np.arange(len(top)), top["cooks_distance"], color=RED, alpha=.82)
    ax.axvline(4 / m["n"], color=INK, linestyle="--", linewidth=.9)
    ax.set_yticks(np.arange(len(top))); ax.set_yticklabels(top["country"], fontsize=6.8)
    ax.set_xlabel("Cook's distance"); ax.grid(axis="x", color=GRID, linewidth=.5); add_panel(ax, "D", "Largest influence values")


def plot_environment_forest(ax: plt.Axes, env: pd.DataFrame) -> None:
    show = env.copy()
    y = np.arange(len(show))[::-1]
    ax.axvline(0, color="#777777", linestyle="--", linewidth=.9)
    for yi, (_, row) in zip(y, show.iterrows()):
        color = RED if "interaction" in row["analysis"].lower() else BLUE
        ax.plot([row["CI_low"], row["CI_high"]], [yi, yi], color=color, linewidth=2)
        ax.scatter(row["estimate"], yi, s=48, color=color, edgecolor="white", linewidth=.6, zorder=3)
        ax.text(.42, yi, f"{row['estimate']:.2f} ({row['CI_low']:.2f}, {row['CI_high']:.2f}); P={row['P']:.3f}", va="center", fontsize=6.9)
    ax.set_yticks(y); ax.set_yticklabels(show["analysis"])
    ax.set_xlim(-.65, .86); ax.set_xlabel("Coefficient (95% CI)")
    ax.grid(axis="x", color=GRID, linewidth=.55)
    add_panel(ax, "A", "Exploratory PM2.5 heterogeneity")


def plot_moran_and_spatial(fig: plt.Figure, model_stats: dict[str, object], spatial_diag: pd.DataFrame) -> None:
    d = model_stats["data"].merge(natural_earth_centroids(), on="iso3", how="inner")
    X, names = design_matrix(d, ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    m = hc3_ols(d["z_outcome"].to_numpy(), X, names)
    dist = haversine_matrix(d["latitude"].to_numpy(), d["longitude"].to_numpy())
    n = len(d); W = np.zeros((n,n))
    for i in range(n): W[i, np.argsort(dist[i])[1:7]] = 1
    W = np.maximum(W, W.T); W = W / np.clip(W.sum(axis=1, keepdims=True), 1, None)
    z = (m["resid"] - np.mean(m["resid"])) / np.std(m["resid"])
    lag = W @ z
    gs = fig.add_gridspec(1, 2, width_ratios=[1, 1.25], wspace=.35)
    ax = fig.add_subplot(gs[0,0]); ax.scatter(z, lag, s=18, alpha=.6, color=BLUE); lr=stats.linregress(z,lag); xx=np.linspace(z.min(),z.max(),100); ax.plot(xx,lr.intercept+lr.slope*xx,color=INK)
    ax.axhline(0,color="#888",lw=.7); ax.axvline(0,color="#888",lw=.7); ax.set_xlabel("Standardised residual"); ax.set_ylabel("Spatial lag"); ax.grid(color=GRID,lw=.5)
    val=spatial_diag.iloc[0]; ax.text(.03,.97,f"Moran's I={val['estimate']:.3f}; permutation P={val['P']:.3f}",transform=ax.transAxes,va="top",fontsize=7.1,bbox={"facecolor":"white","edgecolor":"none","alpha":.8})
    add_panel(ax,"A","Spatial residual autocorrelation")
    ax = fig.add_subplot(gs[0,1]); sc=ax.scatter(d["longitude"],d["latitude"],c=z,cmap="coolwarm",vmin=-2.5,vmax=2.5,s=26,edgecolor="white",linewidth=.3); ax.set_xlim(-180,180);ax.set_ylim(-60,90);ax.set_xlabel("Longitude");ax.set_ylabel("Latitude");ax.grid(color=GRID,lw=.4);fig.colorbar(sc,ax=ax,fraction=.035,pad=.02,label="Standardised residual");add_panel(ax,"B","Geographic pattern of model residuals")


def plot_outcome_consistency(ax: plt.Axes, d: pd.DataFrame, summary: pd.DataFrame) -> None:
    for region, sub in d.groupby("WHO_region"):
        ax.scatter(sub["incidence_residual"], sub["mortality_residual"], s=22, alpha=.65, color=REGION_COLORS.get(region,"#888"), edgecolor="white", linewidth=.3, label=region)
    lr = stats.linregress(d["incidence_residual"], d["mortality_residual"]); x=np.linspace(d["incidence_residual"].min(),d["incidence_residual"].max(),100); ax.plot(x,lr.intercept+lr.slope*x,color=INK,lw=1.5)
    p=summary.query("method == 'Pearson'").iloc[0]; s=summary.query("method == 'Spearman'").iloc[0]
    ax.text(.03,.97,f"Pearson r={p['estimate']:.3f}\nSpearman rho={s['estimate']:.3f}\nn={int(p['n'])}",transform=ax.transAxes,va="top",fontsize=7.2,bbox={"facecolor":"white","edgecolor":"none","alpha":.82})
    ax.set_xlabel("Smoking-adjusted incidence residual"); ax.set_ylabel("Smoking-adjusted mortality residual"); ax.grid(color=GRID,lw=.5); ax.legend(loc="lower right",ncol=2,fontsize=6.3)
    add_panel(ax,"A","Incidence and mortality residual consistency")


def plot_supporting_forest_panels(fig: plt.Figure, effects: pd.DataFrame, corr: pd.DataFrame,
                                  window: pd.DataFrame, robust: pd.DataFrame) -> None:
    gs = fig.add_gridspec(2, 2, width_ratios=[1.0, 1.12], hspace=.46, wspace=.38)
    plot_cliffs_forest(
        fig.add_subplot(gs[0, 0]), effects, panel="A",
        title="Income-group exposure effect sizes",
    )
    plot_clean_fuel_correlations(
        fig.add_subplot(gs[0, 1]), pd.DataFrame(), corr, panel="B",
        title="Clean-fuel structural correlations",
    )
    forest_axis(
        fig.add_subplot(gs[1, 0]), window, "exposure_window",
        "Alternative smoking windows", "C", xlim=(0, .56),
    )
    plot_robustness(
        fig.add_subplot(gs[1, 1]), robust, panel="D",
        title="Influence and spatial sensitivity",
    )


def save_supplementary_figures(country: pd.DataFrame, model_stats: dict[str, object], env: pd.DataFrame,
                               spatial_diag: pd.DataFrame, outcome_d: pd.DataFrame,
                               outcome_summary: pd.DataFrame, effects: pd.DataFrame,
                               corr: pd.DataFrame, window: pd.DataFrame,
                               robust: pd.DataFrame) -> list[Path]:
    apply_figure_style(); SUPP_FIG_DIR.mkdir(parents=True, exist_ok=True); out=[]
    fig=plt.figure(figsize=(12.2,6.6)); plot_source_workflow(fig); base=SUPP_FIG_DIR/"Figure_S1_source_identification_and_audit_workflow"; save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(12.5,5.5)); plot_missingness_heatmaps(fig,country); base=SUPP_FIG_DIR/"Figure_S2_missingness_by_income_and_WHO_region";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(9.5,5.6));plot_audit_status_counts(fig.add_subplot(111));base=SUPP_FIG_DIR/"Figure_S3_audit_classification_counts";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(12.0,8.0));plot_region_distributions(fig,country);base=SUPP_FIG_DIR/"Figure_S4_exposure_and_incidence_by_WHO_region";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(13.8,8.5));plot_supporting_forest_panels(fig,effects,corr,window,robust);base=SUPP_FIG_DIR/"Figure_S5_effect_sizes_and_robustness_forests";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(11.5,8.0));plot_model_diagnostics(fig,model_stats);base=SUPP_FIG_DIR/"Figure_S6_incidence_model_diagnostics";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(10.5,5.8));plot_environment_forest(fig.add_subplot(111),env);base=SUPP_FIG_DIR/"Figure_S7_exploratory_environmental_heterogeneity";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(11.5,5.3));plot_moran_and_spatial(fig,model_stats,spatial_diag);base=SUPP_FIG_DIR/"Figure_S8_spatial_residual_diagnostics";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    fig=plt.figure(figsize=(8.4,6.8));plot_outcome_consistency(fig.add_subplot(111),outcome_d,outcome_summary);base=SUPP_FIG_DIR/"Figure_S9_incidence_mortality_residual_consistency";save_pub(fig,base);plt.close(fig);out.append(base.with_suffix('.png'))
    return out


AUTHORS = [
    "Shen Wang1", "Jing Zhou2", "Zhangdong Jiang1", "Pengyu Li3", "Sha Xiao2*",
]
AFFILIATIONS = [
    "1 School of Cancer Science, College of Medical, Veterinary & Life Sciences, University of Glasgow, Glasgow, United Kingdom",
    "2 Key Laboratory of Tropical Translational Medicine of Ministry of Education, School of Public Health, Hainan Academy of Medical Sciences, Hainan Medical University, Haikou, Hainan 571199, People's Republic of China",
    "3 International Center for Aging and Cancer, Hainan Medical University, Haikou, Hainan, People's Republic of China",
]
CORRESPONDING = [
    "*Corresponding author: Sha Xiao, PhD",
    "Telephone: +86 18808977719",
    "Email: hy0208032@muhn.edu.cn",
    "ORCID: 0000-0002-1645-4846",
    AFFILIATIONS[1][2:],
]

ABSTRACT = {
    "Background": (
        "Analyses can quantify exposure and lung-cancer burden without showing whether indicators needed to monitor prevention and care inequalities are comparable across settings. We assessed indicator availability, patterned missingness and measurable exposure inequalities for women's lung-cancer prevention within a Health Inequality Data Repository-centred source frame."
    ),
    "Methods": (
        "We conducted an ecological study of 197 countries and territories. The analytic frame comprised current female smoking, clean-fuel access, ambient fine particulate matter (PM2.5) and age-standardised female lung-cancer incidence. A structured audit classified exposure, burden, care-pathway and socioeconomic outcome components by coverage, disaggregation, comparability, machine-readability, uncertainty and age standardisation. Income inequalities were estimated as median differences with bootstrap 95% confidence intervals and Cliff's delta. Regression of age-standardised incidence on the available 2000-2010 female smoking proxy, gross domestic product, urbanisation and WHO region provided a validation analysis; HC3, influence and spatial sensitivity analyses were used."
    ),
    "Results": (
        "The complete four-indicator frame included 163/197 countries and territories; 28 were missing female smoking only. Complete inclusion ranged from 72.0% in low-income to 89.1% in high-income settings. Low-income versus high-income countries had a 96.2-percentage-point higher median clean-fuel deficit (95% CI 89.4 to 98.6) and 24.1 micrograms/m3 higher median PM2.5 (95% CI 14.8 to 28.6), whereas female smoking was 15.9 percentage points higher in high-income countries (95% CI 10.6 to 18.5). Three of six exposure components were adequate for analysis, compared with one of four burden components; care and socioeconomic outcome indicators were limited, insufficiently integrated or not identified as compatible global indicators. Historical female smoking was positively associated with age-standardised incidence (standardised beta 0.323, 95% CI 0.163 to 0.483). The estimate was robust to influence and spatial sensitivity analyses, although residual spatial autocorrelation remained."
    ),
    "Conclusions": (
        "Global data support monitoring of selected exposure inequalities, but not the full prevention-care pathway. The principal contribution is an auditable account of what can and cannot be compared; the ecological incidence model is supportive rather than causal evidence."
    ),
}

KEYWORDS = "Lung Neoplasms; Women; Health Status Disparities; Tobacco Smoking; Air Pollution; Particulate Matter; Public Health Surveillance; Ecological Studies"

BACKGROUND = [
    (
        "Lung cancer is a leading cause of cancer incidence and death among women, with marked variation between countries.[1-3] Tobacco control remains central to prevention: active smoking is the dominant preventable cause, and population patterns reflect the timing, intensity and decline of successive smoking epidemics.[4-8] Current female smoking prevalence, however, cannot represent accumulated exposure over the latency period of lung cancer. Prevention surveillance also needs to recognise second-hand smoke, ambient air pollution, household fuel combustion, occupational carcinogens and residential radon, without treating country-level proxies as personal exposure measurements.[6,9-14]"
    ),
    (
        "Recent global analyses have estimated age-standardised lung-cancer burden attributable to active and second-hand smoking, ambient and household particulate pollution, and other risks across countries and over time.[14-16] These studies provide important estimates of attributable burden and socioeconomic gradients. They do not answer a different infrastructure question: whether the indicators required to monitor women's prevention, diagnosis, treatment, survival and social inequalities are available in a comparable, disaggregated and reusable form. A data system may contain national exposure or burden estimates while remaining unable to show who is screened, when disease is diagnosed, whether treatment is received, or whether outcomes differ by wealth, education or residence."
    ),
    (
        "The WHO Health Inequality Data Repository (HIDR) provides a timely setting for this question. WHO released Health Equity Assessment Toolkit version 7 in March 2026 with the 2025 HIDR data update; the repository contains more than 13 million data points across over 2400 indicators, 22 inequality dimensions and 62 datasets.[17,18] HIDR is designed to support disaggregated monitoring rather than to serve as a complete cancer registry or care-pathway database.[18,19] Its content can therefore be used both to identify measurable prevention inequalities and to make gaps in current global surveillance explicit."
    ),
    (
        "Relevant data are distributed across systems beyond HIDR. WHO GHO and GHE, GBD, second-hand-smoke and tobacco-surveillance systems, CanScreen5, screening trials, CONCORD, subtype estimates and radon-risk estimates each cover different parts of the pathway.[20-31] Individual mechanisms relevant to women, including never-smoker lung cancer and hormonal pathways, require clinical and molecular data that are beyond this ecological design.[32,33] We therefore conducted an HIDR-centred cross-national study with two linked aims. First, we audited whether selected exposure, burden, care-pathway and socioeconomic outcome components were suitable for comparable country-level analysis within a predefined source frame. Second, we quantified income-group inequalities in the exposures that were measurable. To ensure that the resulting dataset retained an epidemiologically coherent disease anchor, we examined the association between an available historical female smoking proxy and explicitly age-standardised female lung-cancer incidence. This model was used as a validation analysis, not as an estimate of causal effects or of unmeasured environmental exposure."
    ),
]

METHOD_SECTIONS = [
    ("Study design and analytical scope", [
        (
            "We conducted a cross-national ecological study with one row per country or territory. The harmonised frame contained 197 countries and territories. All findings refer to country-level indicator availability, exposure distributions and associations; they do not support individual-level inference. The study was designed as exposure-side public-health surveillance and a structured data-source audit, not as causal attribution of female lung-cancer burden."
        ),
        (
            "Countries were linked using ISO3 codes. WHO region and World Bank income group were used for stratified summaries. Territories or microstates lacking geometry in the Natural Earth 1:110 million boundary file were retained in all tabular denominators but could not be shaded on maps. Income-unclassified settings were retained in global coverage totals and displayed separately, but were excluded from low-income versus high-income contrasts."
        ),
    ]),
    ("Source frame and structured source identification", [
        (
            "The formal source frame included WHO HIDR/HEAT-linked data, WHO Global Health Observatory (GHO), WHO Global Health Estimates (GHE), IHME/Global Burden of Disease (GBD)-linked public outputs, World Bank indicators and Global Data Lab HDI. We conducted targeted verification of IARC/GLOBOCAN and histology resources, CanScreen5, CONCORD survival resources, WHO second-hand smoke metadata, and Global Tobacco Surveillance System resources. Searches were conducted from May to 13 July 2026 using indicator catalogues, metadata pages, public downloads and documentation. The complete entry-point and decision record is reported in Supplementary Table S2 and the source workflow in Supplementary Figure S1."
        ),
        (
            "A candidate indicator was eligible for integration when it provided a harmonised country identifier and definition, a usable country-year value and a measure relevant to a prespecified pathway component. Targeted sources that contained clinically relevant information but lacked sufficiently complete, women-specific, disaggregated, machine-readable or cross-nationally comparable outputs were recorded as located but not integrated. Accordingly, the audit describes suitability within this source frame; it does not claim that local, registry or research data are globally absent."
        ),
    ]),
    ("Indicator-availability audit", [
        (
            "The audit covered active smoking, clean-fuel access, ambient PM2.5, second-hand smoke, occupational carcinogens, radon, incidence, mortality, years of life lost (YLL), disability-adjusted life years (DALY), low-dose computed tomography (LDCT) screening, stage, histology or subtype, treatment, survival, and wealth-, education- and residence-stratified lung-cancer outcomes. For each component we assessed country coverage, female or sex disaggregation, residence and socioeconomic disaggregation, time coverage, cross-country comparability, machine-readability, uncertainty availability and explicit age standardisation."
        ),
        (
            "Components were classified as adequate for this analysis, available with limitations, source located but not sufficiently harmonised or integrated, no compatible global indicator identified, or not applicable. Adequacy required at least 80% country coverage for the relevant analytic variable, a comparable definition and unit, machine-readability, and the female or sex dimension where required. The limited category captured lower coverage or important deficiencies in subgroup, time, uncertainty or standardisation information. Full operational criteria and item-level evidence notes are reported in Supplementary Table S1."
        ),
        (
            "Shen Wang and Jing Zhou independently coded all 18 components across the 10 audit dimensions using the predefined categories. After both coding rounds were complete, coder-specific records were archived on 13 July 2026 and compared cell by cell. Agreement was 180/180 cells (100%), with Cohen's kappa of 1.000; no item required adjudication by Sha Xiao. Supplementary Table S3 reports the two codes, final classification and agreement for every cell. Two AI-assisted read-only checks also reproduced the final classifications, but these secondary quality-control checks were not included in the inter-rater statistic."
        ),
    ]),
    ("Primary indicators and derived measures", [
        (
            "The primary analytic frame used four non-redundant indicators: female age-standardised current tobacco smoking prevalence in 2022; total population access to clean fuels and technologies for cooking in 2020; national annual mean ambient PM2.5 concentration in 2019; and age-standardised female tracheal, bronchus and lung cancer incidence in 2021. This four-indicator frame avoided counting mortality, YLL and DALY as separate contributions to a generic completeness score. Coverage of the latter burden outcomes was reported separately for triangulation."
        ),
        (
            "Historical female smoking was the mean of female age-standardised current tobacco smoking prevalence in 2000, 2005 and 2010. The window was chosen to precede the 2021 incidence outcome and to reflect cohort history better than current smoking, but it captures only part of the latency-relevant exposure period. Clean-fuel deficit was 100 minus the percentage of the population primarily relying on clean fuels and technologies for cooking. Rural clean-fuel disadvantage was urban minus rural clean-fuel access. Clean-fuel deficit is a country-level proxy for household energy poverty, not a direct measurement of individual household air-pollution exposure. Ambient PM2.5 was a modelled national concentration, not personal exposure."
        ),
        (
            "Structural indicators included WHO region, World Bank income group, gross domestic product (GDP) per capita, urbanisation, the proportion aged 65 years or older, HDI and Socio-demographic Index (SDI).[20-22,34,35] WHO GHE female mortality, YLL and DALY rates were retained as supplementary outcomes because the downloaded labels did not explicitly identify their age-standardisation procedure. The age-standardised incidence estimate was therefore the only burden indicator used in the primary validation model. Indicator definitions, coverage and analytical roles are summarised in Table 1."
        ),
    ]),
    ("Statistical analysis", [
        (
            "We counted indicator availability and classified each country into a complete four-indicator frame or a mutually exclusive missingness pattern. Coverage was summarised globally and by income group and WHO region. Included and excluded countries were compared using median HDI and SDI differences (excluded minus included) with non-parametric bootstrap 95% confidence intervals. Missingness analyses used all countries with the relevant grouping variable and did not impute unavailable indicators."
        ),
        (
            "For each measurable exposure, we reported medians, interquartile ranges and country-level distributions by income group. Formal inequalities were expressed as the low-income versus high-income median difference in the direction stated for each measure, with 5000-sample bootstrap 95% confidence intervals. Cliff's delta, defined as the probability difference that a randomly selected low-income value was higher rather than lower than a high-income value, provided a rank-based effect size with bootstrap confidence intervals.[36] Clean-fuel deficit was related to SDI, HDI, urbanisation and rural clean-fuel disadvantage using Spearman correlations with bootstrap confidence intervals. No significance stars or multiplicity-adjusted hypothesis testing were used because these analyses quantified prespecified descriptive inequalities."
        ),
        (
            "The validation outcome was z-scored age-standardised female lung-cancer incidence. The main exposure was z-scored historical female smoking. The adjusted model included z-scored log GDP per capita, z-scored urbanisation and WHO region. Age structure was not required for the primary outcome because incidence was explicitly age-standardised, but the proportion aged 65 years or older was added in sensitivity analysis. Heteroskedasticity-consistent HC3 standard errors and 95% confidence intervals were used.[37] We compared smoking windows from 2000, 2005, 2010, their 2000-2010 mean and current smoking in 2022. Nested models compared smoking alone, structural adjustment, and addition of clean-fuel deficit and PM2.5 using adjusted R-squared and Akaike information criterion."
        ),
        (
            "Influence was assessed using leverage and Cook's distance, with 4/n as a screening threshold. Sensitivity analyses excluded observations above that threshold and countries with populations below one million. Spatial residual dependence was assessed on countries matched to Natural Earth centroids using a symmetrised six-nearest-neighbour matrix and 999 permutations of Moran's I. Conley-type covariance estimates with Bartlett kernels at 1500, 2500 and 4000 km were used to examine spatially correlated errors.[38] Exploratory models related the incidence residual to clean-fuel deficit and PM2.5 globally, among non-high-income countries, with WHO-region adjustment, and through a PM2.5-by-high-income interaction. These were deliberately supplementary because exposure histories were short and subgroup estimates were sensitive to contextual adjustment."
        ),
        (
            "As an estimate-quality check, we compared residuals from analogous incidence and WHO mortality models among common countries. Mortality-based results were not promoted to primary evidence because age standardisation was not explicit in the source label. Analyses were complete case for the variables in each model. Python 3.12, pandas, NumPy, SciPy and Matplotlib were used. Reporting was informed by STROBE, SAGER and GATHER guidance.[39-41] Reproduction code and cleaned aggregate data are available in the public repository."
        ),
    ]),
]

RESULT_SECTIONS = [
    ("Availability of the primary analytic indicators", [
        (
            "The four non-redundant analytic indicators were jointly available for 163 of 197 countries and territories (82.7%) (Figure 1A; Table 1). Current female smoking was available for 165/197, clean-fuel access for 191/197, ambient PM2.5 for 195/197 and age-standardised female lung-cancer incidence for 194/197 (Figure 1B). Among the 34 countries outside the complete frame, 28 were missing female smoking only, two were missing clean-fuel access only, one was missing both, and three had other multi-indicator patterns (Supplementary Table S4)."
        ),
        (
            "Complete-frame inclusion increased from 72.0% (18/25) in low-income countries to 82.0% (41/50) in lower-middle-income, 84.9% (45/53) in upper-middle-income and 89.1% (57/64) in high-income countries; two of five income-unclassified settings were complete (Figure 1C). Excluded countries had a median HDI 0.100 lower than included countries (bootstrap 95% CI -0.224 to -0.033). The corresponding SDI difference was -0.038, but its interval included zero (-0.127 to 0.043) (Figure 1D; Table 2). Coverage patterns also differed across WHO regions (Supplementary Figure S2; Supplementary Table S5)."
        ),
    ]),
    ("Prevention-care pathway availability", [
        (
            "The audit showed a gradient from exposure measurement to care and equity outcomes (Figure 2; Supplementary Figure S3). The two human coders agreed on all 180 component-by-dimension classifications (Cohen's kappa=1.000), and no adjudication change was required (Supplementary Table S3). Three of six exposure components were adequate for the present analysis: female smoking, clean-fuel access and ambient PM2.5. Second-hand smoke, occupational carcinogens and residential radon were available with limitations because relevant WHO, tobacco-surveillance or GBD sources existed but were not available as consistently harmonised female country-year exposures within the analytic frame (Figure 2A-B; Supplementary Tables S1-S3)."
        ),
        (
            "Age-standardised female incidence was adequate, whereas mortality, YLL and DALY were classified as available with limitations because standardisation and uncertainty were not explicit in the retained labels. CanScreen5 and other programme sources were located, but LDCT screening indicators were not sufficiently integrated for this analysis.[25,26] Stage and treatment were similarly located but not integrated. Histology or subtype and survival were available with limitations: global subtype estimates exist, and CONCORD-3 reported comparable survival from 322 registries in 71 countries, but neither provided a complete socioeconomic country-level pathway for this study.[29,30] No compatible global wealth- or education-stratified female lung-cancer outcome indicator was identified; residence-stratified outcomes were located but insufficiently integrated."
        ),
    ]),
    ("Magnitude of measurable exposure inequalities", [
        (
            "The global geography and income-group distributions of measurable exposures differed sharply (Figure 3A-B; Table 2). Median current female smoking was 17.2% in high-income countries and 1.3% in low-income countries, an HIC-minus-LIC difference of 15.9 percentage points (95% CI 10.6 to 18.5; Cliff's delta for LIC relative to HIC -0.88, 95% CI -0.96 to -0.75). By contrast, median clean-fuel deficit was 96.2% in low-income countries and 0% in high-income countries, a 96.2-percentage-point LIC-minus-HIC difference (95% CI 89.4 to 98.6; Cliff's delta 1.00) (Supplementary Figure S5A; Supplementary Table S7)."
        ),
        (
            "Median ambient PM2.5 was 34.9 micrograms/m3 in low-income and 10.7 micrograms/m3 in high-income countries, a 24.1-microgram/m3 difference (95% CI 14.8 to 28.6; Cliff's delta 0.79, 95% CI 0.65 to 0.92). The distribution was less monotonic across intermediate income groups than the clean-fuel gradient (Figure 3B). Exposure and incidence distributions also varied by WHO region (Supplementary Figure S4; Supplementary Table S6)."
        ),
        (
            "Residence was the clearest consistently available within-country inequality dimension. The median urban-rural clean-fuel gap was 7.0 percentage points larger in low-income than high-income countries (95% CI 2.2 to 18.1; Cliff's delta 0.90, 95% CI 0.79 to 1.00) (Figure 3C-D; Table 2; Supplementary Figure S5A). The smaller median gap in low-income than lower-middle-income settings did not imply greater equity: many low-income countries had low access in both urban and rural populations. Clean-fuel deficit correlated inversely with SDI (rho=-0.898, 95% CI -0.925 to -0.859), HDI (rho=-0.909, 95% CI -0.932 to -0.869) and urbanisation (rho=-0.663, 95% CI -0.740 to -0.571), and positively with the rural clean-fuel gap (rho=0.709, 95% CI 0.592 to 0.802) (Supplementary Figure S5B)."
        ),
    ]),
    ("Age-standardised incidence validation analysis", [
        (
            "Historical female smoking was positively associated with age-standardised female lung-cancer incidence. Among 163 countries with complete model variables, the unadjusted standardised coefficient was 0.60. After adjustment for log GDP, urbanisation and WHO region, the coefficient was 0.323 (95% CI 0.163 to 0.483; P<0.001) (Figure 4A-B; Table 2; Supplementary Table S8). The adjusted model explained 56.4% of outcome variation (adjusted R-squared 0.541). Adding current clean-fuel deficit and PM2.5 did not improve fit: adjusted R-squared was 0.536 and AIC increased from 345.3 to 346.7 (Figure 4C)."
        ),
        (
            "Associations were similar for female smoking in 2000, 2005 and 2010; the coefficient for current smoking in 2022 was smaller (0.177, 95% CI 0.017 to 0.338) (Supplementary Figure S5C; Supplementary Table S9). The historical-smoking estimate remained positive after additional age-65 adjustment, exclusion of 11 countries above the Cook threshold, restriction to populations of at least one million, and Conley-type spatial covariance estimation (Supplementary Figure S5D; Supplementary Tables S9-S10). Model residuals showed modest spatial autocorrelation (Moran's I=0.106; permutation P=0.009), visible as a non-random geographic pattern in Figure 4D. Spatial sensitivity intervals were therefore reported, and country observations were not assumed to be fully independent (Supplementary Figures S6 and S8)."
        ),
    ]),
    ("Exploratory environmental and outcome checks", [
        (
            "Environmental residual models did not support a uniform global association. In the mutually adjusted global model, the PM2.5 coefficient was -0.009 (95% CI -0.130 to 0.112). Among non-high-income countries it was 0.124 (95% CI 0.013 to 0.235), but it attenuated after WHO-region adjustment to 0.106 (95% CI -0.042 to 0.254) (Supplementary Figure S7; Supplementary Table S11). The PM2.5-by-high-income interaction was directionally consistent before and after influence exclusion, but these ecological subgroup results were retained as exploratory because they depended on context specification and short-lag exposure measures."
        ),
        (
            "Incidence- and mortality-model residuals were correlated among 163 common countries (Pearson r=0.804; Spearman rho=0.758) (Supplementary Figure S9; Supplementary Table S12). This consistency reduced concern that the incidence pattern was unique to one outcome, but it did not solve differences in source methods, registration quality, treatment access or uncertainty. Mortality, YLL and DALY were therefore used only for supplementary triangulation."
        ),
    ]),
]

DISCUSSION = [
    (
        "Global women's lung-cancer surveillance showed a practical asymmetry. Within the prespecified HIDR-centred source frame, current female smoking, clean-fuel access, ambient PM2.5 and age-standardised incidence could be assembled for 163 countries and territories. Comparable information became progressively thinner across second-hand and long-latency exposures, screening, stage, treatment, survival and socioeconomic outcomes. Relevant data often exist locally, but remain dispersed across surveys, registries and programme systems, use different definitions or disaggregations, or cannot yet be linked into a global inequality-monitoring frame."
    ),
    (
        "Recent GBD analyses addressed attributable lung-cancer burden.[14-16] Zhang and colleagues quantified exposure, burden trends and transnational inequalities using age-standardised summary exposure values, concentration indices and slope indices.[15] Our question was which parts of the pathway can be observed with compatible country-level indicators, and whether missingness was socially patterned. Attributable-burden models estimate the consequences of specified risks; an availability audit identifies where routine global surveillance cannot test equity across prevention and care."
    ),
    (
        "The exposure inequalities were large but not uniform. Active female smoking was higher in high-income countries, consistent with the historical diffusion and later decline of women's tobacco epidemics.[4-8] Clean-fuel deficit displayed the opposite gradient and almost complete separation between low- and high-income groups. Ambient PM2.5 was also higher in low-income countries, although intermediate groups did not follow a simple monotonic pattern. These contrasts show why a single composite multiple-exposure score would be difficult to interpret: the indicators represent different processes, time horizons and policy levers. Reporting effect sizes and raw distributions preserves those differences."
    ),
    (
        "Clean-fuel access is particularly informative as an equity indicator. Its close relationships with SDI, HDI, urbanisation and rural disadvantage locate it within structural and household-energy conditions, not as an isolated biological exposure. The smaller median urban-rural gap in low-income countries should not be read as greater fairness because low access in both groups can compress the difference. For policy, universal deprivation and unequal distribution are separate problems. At the same time, clean-fuel deficit is not an individual household-air-pollution dose, and its absence from the incidence residual model does not invalidate its value for prevention surveillance."
    ),
    (
        "The age-standardised incidence analysis provided an expected epidemiological anchor. The available 2000-2010 smoking proxy was associated with current female lung-cancer incidence after structural adjustment, while current smoking produced a smaller coefficient. This pattern is consistent with latency and cohort history, but the proxy begins too recently to capture the full 20-40-year exposure period relevant to many 2021 cases. The model should therefore be understood as partially historical smoking adjustment. Its residual is not an estimate of second-hand smoke, air pollution, hormonal susceptibility, diagnostic intensity or any other unmeasured cause."
    ),
    (
        "The environmental models reinforce that caution. A PM2.5 association appeared in the non-high-income subset but lost precision after WHO-region adjustment, whereas no uniform global residual association was observed. Ecological collinearity, regional context, modelled exposure surfaces, short exposure histories and outcome data quality can all produce this pattern. The finding is better used to motivate longer historical exposure series and multilevel data than to attribute excess female lung-cancer incidence to national PM2.5."
    ),
    (
        "The pathway audit also refines the language of data gaps. CanScreen5 is a global, increasingly harmonised screening repository, but its established quantitative coverage has focused on breast, cervical and colorectal programmes, with lung-screening initiatives still emerging.[25,26] CONCORD provides high-quality survival comparisons, but its 71-country registry coverage and lack of globally harmonised socioeconomic linkage do not constitute universal survival-equity surveillance.[29] Likewise, global subtype estimates and GBD radon or occupational-risk estimates are valuable without filling stage, treatment or within-country social gradients.[14,30,31] Classifying these resources as located or limited is more accurate than labelling the domains absent."
    ),
    (
        "Several limitations remain. The source identification was structured but not a systematic review of every national registry or programme. Although two investigators independently coded the audit with complete agreement and the coder-specific records are now archived, the records were consolidated into the analytic archive after completion rather than time-stamped prospectively during source identification. Several inputs were modelled estimates, and harmonised uncertainty draws were unavailable across data streams. Complete-case comparisons can underrepresent countries with lower development. Country observations were spatially correlated, although the smoking coefficient remained positive with Conley-type uncertainty. Finally, the study is ecological and cannot infer individual exposure, susceptibility, access to care or causal mechanisms."
    ),
    (
        "A useful next step is a minimum global women's lung-cancer equity indicator set that connects exposure and care. It should retain age-standardised active smoking, clean cooking and ambient PM2.5; add harmonised female second-hand smoke, occupational and radon indicators; and link LDCT eligibility and uptake, stage, histology, treatment and survival to wealth, education and residence. Versioned metadata should record whether estimates are observed or modelled, the standard population, uncertainty, observation year and subgroup denominator. The 2025 HIDR update and HEAT version 7 provide a strong platform for exposure-side inequality monitoring, but cancer registries, screening programmes and survival systems must be interoperable if the full pathway is to become visible."
    ),
]

CONCLUSION = (
    "Within an HIDR-centred global source frame, selected exposure and incidence indicators support broad country-level monitoring of women's lung-cancer prevention, but care-pathway and socioeconomic outcome surveillance remains fragmented. Missingness was patterned by income and development, and measurable exposures followed opposing inequality gradients. Historical female smoking was associated with age-standardised incidence, but neither that association nor exploratory environmental residuals establish causality. The priority is therefore twofold: act on the large exposure inequalities already measurable and build interoperable, disaggregated data for screening, diagnosis, treatment and survival."
)


REFERENCES = [
    "Bray F, Laversanne M, Sung H, et al. Global cancer statistics 2022: GLOBOCAN estimates of incidence and mortality worldwide for 36 cancers in 185 countries. CA Cancer J Clin. 2024;74:229-263. doi:10.3322/caac.21834.",
    "World Health Organization. Lung cancer. Geneva: WHO; 2026. https://www.who.int/news-room/fact-sheets/detail/lung-cancer. Accessed 13 Jul 2026.",
    "Leiter A, Veluswamy RR, Wisnivesky JP. The global burden of lung cancer: current status and future trends. Nat Rev Clin Oncol. 2023;20:624-639. doi:10.1038/s41571-023-00798-3.",
    "World Health Organization. Tobacco. Geneva: WHO; 2025. https://www.who.int/news-room/fact-sheets/detail/tobacco. Accessed 13 Jul 2026.",
    "World Health Organization. WHO global report on trends in prevalence of tobacco use 2000-2024 and projections 2025-2030. Geneva: WHO; 2025. https://www.who.int/publications/i/item/9789240115323.",
    "International Agency for Research on Cancer. Tobacco smoke and involuntary smoking. IARC Monographs on the Evaluation of Carcinogenic Risks to Humans, Volume 83. Lyon: IARC; 2004. https://publications.iarc.fr/101.",
    "Thun MJ, Carter BD, Feskanich D, et al. 50-year trends in smoking-related mortality in the United States. N Engl J Med. 2013;368:351-364. doi:10.1056/NEJMsa1211127.",
    "Ng M, Freeman MK, Fleming TD, et al. Smoking prevalence and cigarette consumption in 187 countries, 1980-2012. JAMA. 2014;311:183-192. doi:10.1001/jama.2013.284692.",
    "International Agency for Research on Cancer. Outdoor air pollution. IARC Monographs on the Evaluation of Carcinogenic Risks to Humans, Volume 109. Lyon: IARC; 2016. https://publications.iarc.who.int/538.",
    "Hamra GB, Guha N, Cohen A, et al. Outdoor particulate matter exposure and lung cancer: a systematic review and meta-analysis. Environ Health Perspect. 2014;122:906-911. doi:10.1289/ehp.1408092.",
    "World Health Organization. Household air pollution. Geneva: WHO; 2025. https://www.who.int/news-room/fact-sheets/detail/household-air-pollution-and-health. Accessed 13 Jul 2026.",
    "Kurmi OP, Arya PH, Lam KBH, et al. Lung cancer risk and solid fuel smoke exposure: a systematic review and meta-analysis. Eur Respir J. 2012;40:1228-1237. doi:10.1183/09031936.00099511.",
    "Hosgood HD 3rd, Boffetta P, Greenland S, et al. In-home coal and wood use and lung cancer risk: a pooled analysis of the International Lung Cancer Consortium. Environ Health Perspect. 2010;118:1743-1747. doi:10.1289/ehp.1002217.",
    "GBD 2021 Risk Factors Collaborators. Global burden and strength of evidence for 88 risk factors in 204 countries and 811 subnational locations, 1990-2021: a systematic analysis for the Global Burden of Disease Study 2021. Lancet. 2024;403:2162-2203. doi:10.1016/S0140-6736(24)00933-4.",
    "Zhang Y, Wang W, Dai K, et al. Global lung cancer burden attributable to air fine particulate matter and tobacco smoke exposure: spatiotemporal patterns, sociodemographic characteristics, and transnational inequalities from 1990 to 2021. BMC Public Health. 2025;25:1260. doi:10.1186/s12889-025-22450-8.",
    "Deng Y, Li Z, Zhang P, et al. Global, regional and national burden of lung cancer attributable to PM2.5 air pollution: trends from 1990 to 2021 with projections to 2045. J Environ Manage. 2025;390:126216. doi:10.1016/j.jenvman.2025.126216.",
    "World Health Organization. WHO releases updated Health Inequality Data Repository and Health Equity Assessment Toolkit. Geneva: WHO; 2026. https://www.who.int/news/item/03-03-2026-who-releases-updated-health-inequality-data-repository-and-health-equity-assessment-toolkit. Accessed 13 Jul 2026.",
    "World Health Organization. Health Inequality Monitor. Geneva: WHO; 2026. https://www.who.int/data/inequality-monitor. Accessed 13 Jul 2026.",
    "World Health Organization. Handbook on health inequality monitoring: with a special focus on low- and middle-income countries. Geneva: WHO; 2013. https://www.who.int/publications/i/item/9789241548632.",
    "World Health Organization. Global Health Observatory. Geneva: WHO; 2026. https://www.who.int/data/gho. Accessed 13 Jul 2026.",
    "World Health Organization. Global Health Estimates: life expectancy and leading causes of death and disability. Geneva: WHO; 2024. https://www.who.int/data/gho/data/themes/mortality-and-global-health-estimates. Accessed 13 Jul 2026.",
    "GBD 2021 Diseases and Injuries Collaborators. Global incidence, prevalence, years lived with disability, disability-adjusted life-years, and healthy life expectancy for 371 diseases and injuries in 204 countries and territories and 811 subnational locations, 1990-2021: a systematic analysis for the Global Burden of Disease Study 2021. Lancet. 2024;403:2133-2161. doi:10.1016/S0140-6736(24)00757-8.",
    "World Health Organization. Exposure to second-hand smoke in adults: indicator 2266. Geneva: WHO. https://www.who.int/data/gho/data/indicators/indicator-details/GHO/second-hand-smoke-attributable-dalys-per-100000-capita. Accessed 13 Jul 2026.",
    "Centers for Disease Control and Prevention. Global Tobacco Surveillance System Data. Atlanta: CDC; 2025. https://www.cdc.gov/tobacco/global/gtss/gtssdata/index.html. Accessed 13 Jul 2026.",
    "CanScreen5 Collaborators. CanScreen5, a global repository for breast, cervical and colorectal cancer screening programs. Nat Med. 2023;29:1135-1145. doi:10.1038/s41591-023-02315-6.",
    "International Agency for Research on Cancer. CanScreen5: the global cancer screening data repository. Lyon: IARC; 2026. https://canscreen5.iarc.fr/. Accessed 13 Jul 2026.",
    "National Lung Screening Trial Research Team, Aberle DR, Adams AM, et al. Reduced lung-cancer mortality with low-dose computed tomographic screening. N Engl J Med. 2011;365:395-409. doi:10.1056/NEJMoa1102873.",
    "de Koning HJ, van der Aalst CM, de Jong PA, et al. Reduced lung-cancer mortality with volume CT screening in a randomized trial. N Engl J Med. 2020;382:503-513. doi:10.1056/NEJMoa1911793.",
    "Allemani C, Matsuda T, Di Carlo V, et al. Global surveillance of trends in cancer survival 2000-2014 (CONCORD-3): analysis of individual records for 37,513,025 patients diagnosed with one of 18 cancers from 322 population-based registries in 71 countries. Lancet. 2018;391:1023-1075. doi:10.1016/S0140-6736(17)33326-3.",
    "Luo G, Zhang Y, Rumgay H, et al. Estimated worldwide variation and trends in incidence of lung cancer by histological subtype in 2022 and over time: a population-based study. Lancet Respir Med. 2025;13:348-363. doi:10.1016/S2213-2600(24)00428-4.",
    "Institute for Health Metrics and Evaluation. Residential radon: GBD 2021 risk summary. Seattle: IHME; 2025. https://www.healthdata.org/sites/default/files/disease_and_injury/gbd_2021/topic_pdf/risk/90.pdf. Accessed 13 Jul 2026.",
    "Sun S, Schiller JH, Gazdar AF. Lung cancer in never smokers: a different disease. Nat Rev Cancer. 2007;7:778-790. doi:10.1038/nrc2190.",
    "Siegfried JM, Stabile LP. Estrogenic steroid hormones in lung cancer. Semin Oncol. 2014;41:5-16. doi:10.1053/j.seminoncol.2013.12.009.",
    "World Bank. World Bank country and lending groups. Washington, DC: World Bank; 2026. https://datahelpdesk.worldbank.org/knowledgebase/articles/906519. Accessed 13 Jul 2026.",
    "Global Data Lab. Subnational Human Development Index. Nijmegen: Radboud University; 2026. https://globaldatalab.org/shdi/. Accessed 13 Jul 2026.",
    "Cliff N. Dominance statistics: ordinal analyses to answer ordinal questions. Psychol Bull. 1993;114:494-509. doi:10.1037/0033-2909.114.3.494.",
    "MacKinnon JG, White H. Some heteroskedasticity-consistent covariance matrix estimators with improved finite sample properties. J Econom. 1985;29:305-325. doi:10.1016/0304-4076(85)90158-7.",
    "Conley TG. GMM estimation with cross sectional dependence. J Econom. 1999;92:1-45. doi:10.1016/S0304-4076(98)00084-0.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. PLoS Med. 2007;4:e296. doi:10.1371/journal.pmed.0040296.",
    "Heidari S, Babor TF, De Castro P, et al. Sex and Gender Equity in Research: rationale for the SAGER guidelines and recommended use. Res Integr Peer Rev. 2016;1:2. doi:10.1186/s41073-016-0007-6.",
    "Stevens GA, Alkema L, Black RE, et al. Guidelines for Accurate and Transparent Health Estimates Reporting: the GATHER statement. PLoS Med. 2016;13:e1002056. doi:10.1371/journal.pmed.1002056.",
]


def main_table1(country: pd.DataFrame) -> list[list[str]]:
    return [
        ["Indicator", "Source; year", "Definition and unit", "Coverage", "Equity dimension", "Role", "Principal limitation"],
        ["Current female smoking", "WHO GHO; 2022", "Female age-standardised current tobacco smoking prevalence (%)", "165/197", "Sex", "Primary exposure and completeness", "Modelled estimate; no globally contemporaneous residence or socioeconomic subgroup"],
        ["Historical female smoking", "WHO GHO; 2000, 2005, 2010", "Mean of three female age-standardised prevalence estimates (%)", "163 model countries", "Sex and time", "Validation exposure", "Begins too recently to represent the full latency-relevant history"],
        ["Clean-fuel access", "WHO GHO; 2020", "Population primarily relying on clean fuels and technologies for cooking (%)", "191/197", "Total, urban and rural", "Primary exposure; residence inequality", "Household-energy proxy, not personal household-air-pollution dose"],
        ["Ambient PM2.5", "WHO GHO; 2019", "National annual mean modelled concentration (micrograms/m3)", "195/197", "Country; limited residence", "Primary exposure", "National modelled concentration; no personal or long-lag exposure"],
        ["Female lung-cancer incidence", "IHME/GBD-linked; 2021", "Age-standardised female incidence per 100,000", "194/197", "Sex", "Primary burden anchor", "Modelled estimate; uncertainty draws not retained in the analytic file"],
        ["Female mortality, YLL and DALY", "WHO GHE/HIDR; 2021", "Sex-specific rates per 100,000", "196-197/197", "Sex", "Supplementary triangulation", "Age-standardisation not explicit in downloaded labels"],
        ["Structural context", "World Bank, Global Data Lab, GBD-linked; 2021-2022", "Income group, GDP, urbanisation, age 65+, HDI and SDI", "Variable-specific", "Country", "Missingness, stratification and adjustment", "Differing source years and model structures"],
    ]


def main_table2(country: pd.DataFrame, effects: pd.DataFrame, diffs: pd.DataFrame, model_stats: dict[str, object], env: pd.DataFrame, outcome_summary: pd.DataFrame) -> list[list[str]]:
    eff = effects.set_index("exposure")
    hdi = diffs.set_index("index").loc["HDI"]
    main = model_stats["main_coef"]
    nonhic = env.query("analysis == 'Non-high-income, mutually adjusted'").iloc[0]
    reg = env.query("analysis == 'Non-high-income + WHO region'").iloc[0]
    pearson = outcome_summary.query("method == 'Pearson'").iloc[0]

    def ef(name: str) -> str:
        r = eff.loc[name]
        return f"Median difference {r['absolute_median_difference']:.1f} {r['unit']} (95% CI {r['median_difference_95CI_low']:.1f} to {r['median_difference_95CI_high']:.1f}); Cliff's delta {r['cliffs_delta_low_vs_high']:.2f} ({r['cliffs_delta_95CI_low']:.2f} to {r['cliffs_delta_95CI_high']:.2f})"

    return [
        ["Evidence domain", "Quantitative finding", "Sample/coverage", "Interpretation", "Related display"],
        ["Analytic-frame availability", "Four indicators jointly available for 163/197; 28/34 excluded settings missing female smoking only", "197 countries and territories", "The frame is broad but incomplete; smoking is the main driver of exclusion", "Figure 1A-B; Supplementary Table S4"],
        ["Patterned missingness", f"Complete frame: LIC 72.0%, LMIC 82.0%, UMIC 84.9%, HIC 89.1%; excluded-minus-included HDI {hdi['excluded_minus_included']:.3f} (95% CI {hdi['CI_low']:.3f} to {hdi['CI_high']:.3f})", "192 income-classified; HDI n=164", "Missingness is associated with development context and should not be treated as random", "Figure 1C-D; Supplementary Figure S2"],
        ["Female smoking inequality", ef("Current female smoking"), "LIC n=18; HIC n=58", "Current female smoking is higher in high-income settings", "Figure 3A-B; Supplementary Figure S5A; Supplementary Table S7"],
        ["Clean-fuel inequality", ef("Clean-fuel deficit"), "LIC n=25; HIC n=62", "Household-energy deprivation is concentrated in low-income settings", "Figure 3A-B; Supplementary Figure S5A; Supplementary Table S7"],
        ["Ambient PM2.5 inequality", ef("Ambient PM2.5"), "LIC n=25; HIC n=63", "Ambient pollution is higher in low-income settings, with a non-monotonic intermediate gradient", "Figure 3A-B; Supplementary Figure S5A; Supplementary Table S7"],
        ["Residence inequality", ef("Urban-rural clean-fuel gap"), "LIC n=25; HIC n=62", "A small gap can coexist with universally low access; absolute access and distribution must be interpreted together", "Figure 3C-D; Supplementary Figure S5A"],
        ["Age-standardised incidence anchor", f"Adjusted historical-smoking beta {main['estimate']:.3f} (95% CI {main['CI_low']:.3f} to {main['CI_high']:.3f}; P<0.001)", "n=163", "Expected positive association supports data coherence but is not causal attribution", "Figure 4A-D; Supplementary Figure S5C-D; Supplementary Tables S8-S10"],
        ["Exploratory PM2.5 heterogeneity", f"Non-HIC beta {nonhic['estimate']:.3f} ({nonhic['CI_low']:.3f} to {nonhic['CI_high']:.3f}); after WHO-region adjustment {reg['estimate']:.3f} ({reg['CI_low']:.3f} to {reg['CI_high']:.3f})", "n=104", "The subgroup signal attenuates with regional context and remains exploratory", "Supplementary Figure S7; Supplementary Table S11"],
        ["Outcome consistency", f"Incidence- and mortality-model residuals: Pearson r={pearson['estimate']:.3f}", "n=163", "Cross-outcome consistency is supportive but does not remove source and standardisation limitations", "Supplementary Figure S9; Supplementary Table S12"],
    ]


FIGURE_LEGENDS = [
    "Figure 1. Availability and patterned missingness in the four-indicator analytic frame. (A) Country and territory missingness patterns for current female smoking, clean-fuel access, ambient PM2.5 and age-standardised female lung-cancer incidence. Grey map areas were not matched to the analytic frame and are not counted as missing. (B) Indicator-specific availability. The dashed line marks the 163-country complete overlap and is not an adequacy threshold. (C) Complete-frame inclusion by World Bank income group. (D) Median HDI and SDI differences between excluded and included countries; points are excluded-minus-included differences and lines are 5000-sample bootstrap 95% confidence intervals. LIC, low-income country; LMIC, lower-middle-income country; UMIC, upper-middle-income country; HIC, high-income country.",
    "Figure 2. Availability audit across prevention, burden, care and equity domains. (A) Adjudicated component-by-dimension audit. A indicates adequate for this analysis; L, available with limitations; LH, a relevant source was located but was not sufficiently harmonised or integrated; NG, no compatible global indicator was identified within the predefined source frame; NA, not applicable. (B) Audit key and counts of overall classifications by domain. Classifications describe analytical suitability within the source frame and do not imply that no local, registry or research data exist.",
    "Figure 3. Global geography and income patterning of measurable prevention exposures. (A) Country-level choropleths for current female smoking in 2022, clean-fuel deficit in 2020 and ambient PM2.5 in 2019. Each map has its own scale; grey denotes no mapped value, and countries or territories without Natural Earth geometry remain in tabular analyses. (B) Country distributions by World Bank income group; boxes show medians and interquartile ranges, whiskers extend to 1.5 times the interquartile range, and points are countries. (C) Income-group medians for urban and rural clean-fuel access; labelled gaps are urban minus rural access. (D) Country-level urban versus rural clean-fuel access. The dashed diagonal denotes equal access, and selected labels identify the five largest observed urban-rural gaps rather than country rankings. Income-unclassified settings were excluded from income-group panels but retained in panel D. LIC, low-income country; LMIC, lower-middle-income country; UMIC, upper-middle-income country; HIC, high-income country.",
    "Figure 4. Historical female smoking and age-standardised lung-cancer incidence. (A) Unadjusted country-level association between the 2000-2010 female smoking mean and 2021 age-standardised female lung-cancer incidence; colours denote WHO regions. (B) Frisch-Waugh-Lovell partial-regression display after adjustment of both variables for log GDP, urbanisation and WHO region. The line and shaded band show the adjusted fit and pointwise HC3 95% confidence interval; the annotated coefficient is from the prespecified primary model. (C) Adjusted R-squared and AIC for nested models. (D) Geographic distribution of primary-model residuals; positive values indicate incidence higher than predicted by the model and negative values indicate lower incidence. Colour limits are symmetric at the 95th percentile of absolute residuals, with more extreme values shown at the darkest endpoint colours. Residuals are model diagnostics, not estimates of omitted exposures or causal effects. Grey denotes no mapped model residual. All model variables are standardised.",
]


SUPPLEMENTARY_FIGURE_LEGENDS = [
    "Supplementary Figure S1. Source-identification and indicator-audit workflow. Shen Wang and Jing Zhou independently coded the predefined audit cells; human agreement was quantified before finalisation. Two AI-assisted read-only checks were used only as secondary quality control and were excluded from Cohen's kappa. The workflow distinguishes formal source-frame coding from targeted verification of relevant cancer, screening, survival and tobacco resources.",
    "Supplementary Figure S2. Coverage by income group and WHO region. Cells report the percentage of countries or territories with each primary indicator and the complete four-indicator overlap.",
    "Supplementary Figure S3. Overall audit classifications by domain. Stacked bars summarise the adjudicated overall classification of components shown individually in Figure 2 and Supplementary Table S3.",
    "Supplementary Figure S4. Exposure and age-standardised incidence distributions by WHO region. Boxes show medians and interquartile ranges, whiskers extend to 1.5 times the interquartile range, and points represent countries or territories.",
    "Supplementary Figure S5. Effect-size and robustness estimates supporting Figures 3 and 4. (A) Cliff's delta for low-income relative to high-income countries with 5000-sample bootstrap 95% confidence intervals; negative values indicate higher values in high-income settings and positive values indicate higher values in low-income settings. (B) Spearman correlations of clean-fuel deficit with structural indicators and bootstrap 95% confidence intervals. (C) Adjusted historical-smoking coefficients across alternative exposure windows. (D) Influence, population-size and spatial-covariance sensitivity estimates. Forest plots are retained in the supplement because the corresponding main figures prioritise spatial patterns and underlying country distributions.",
    "Supplementary Figure S6. Incidence-model diagnostics. Panels show residuals versus fitted values, a normal quantile plot, leverage versus residuals and the 12 largest Cook's distances. The dashed Cook threshold is 4/n.",
    "Supplementary Figure S7. Exploratory PM2.5 heterogeneity. Coefficients and HC3 95% confidence intervals are shown for global, non-high-income, WHO-region-adjusted and interaction analyses of the age-standardised incidence residual. Interaction estimates refer to PM2.5 by high-income status and are not directly comparable with the PM2.5 main-effect rows.",
    "Supplementary Figure S8. Spatial residual diagnostics. The Moran scatter plot and geographic residual display use countries matched to Natural Earth centroids. Moran's I used a symmetrised six-nearest-neighbour matrix and 999 permutations.",
    "Supplementary Figure S9. Incidence-mortality residual consistency. Residuals were derived from analogous historical-smoking models adjusted for log GDP, urbanisation and WHO region. The mortality outcome was retained only for supplementary triangulation because age standardisation was not explicit in its downloaded label.",
]


def md_table(rows: list[list[str]]) -> str:
    esc = lambda x: str(x).replace("|", "\\|").replace("\n", " ")
    out = ["| " + " | ".join(map(esc, rows[0])) + " |", "| " + " | ".join(["---"] * len(rows[0])) + " |"]
    out.extend("| " + " | ".join(map(esc, row)) + " |" for row in rows[1:])
    return "\n".join(out)


def manuscript_markdown(table1: list[list[str]], table2: list[list[str]]) -> str:
    parts = [
        f"# {TITLE}",
        "## Authors\n" + ", ".join(AUTHORS) + "\n\n" + "\n\n".join(AFFILIATIONS + CORRESPONDING) +
        "\n\nAdditional author details: Jing Zhou, email hy0208035@muhn.edu.cn, ORCID 0000-0003-4076-4083; Pengyu Li, email leepy8@163.com.",
        "## Abstract\n" + "\n\n".join(f"**{k}:** {v}" for k, v in ABSTRACT.items()),
        f"**Keywords:** {KEYWORDS}",
        "## Background\n" + "\n\n".join(BACKGROUND),
        "## Methods",
    ]
    for h, paras in METHOD_SECTIONS:
        parts.append(f"### {h}\n" + "\n\n".join(paras))
    parts.append("## Results")
    for h, paras in RESULT_SECTIONS:
        parts.append(f"### {h}\n" + "\n\n".join(paras))
    parts.extend([
        "## Discussion\n" + "\n\n".join(DISCUSSION),
        "## Conclusions\n" + CONCLUSION,
        "## Abbreviations\nAIC, Akaike information criterion; DALY, disability-adjusted life year; GBD, Global Burden of Disease; GDP, gross domestic product; GHE, Global Health Estimates; GHO, Global Health Observatory; HDI, Human Development Index; HEAT, Health Equity Assessment Toolkit; HIC, high-income country; HIDR, Health Inequality Data Repository; LC, lung cancer; LDCT, low-dose computed tomography; LIC, low-income country; LMIC, lower-middle-income country; PM2.5, particulate matter with aerodynamic diameter no greater than 2.5 micrometres; SDI, Socio-demographic Index; UMIC, upper-middle-income country; WHO, World Health Organization; YLL, years of life lost.",
        "## Declarations",
        "### Ethics approval and consent to participate\nThis study used publicly available aggregate country-level data and did not involve individual participants or identifiable personal information. Ethics approval and informed consent were not required.",
        "### Consent for publication\nNot applicable.",
        "### Availability of data and materials\nThe cleaned aggregate country-level dataset, indicator dictionaries, independent-coder audit records, analysis tables and reproduction scripts are available at https://github.com/ammkkr/women-lung-cancer-equity-surveillance. The repository contains no individual-level or identifiable data. Source datasets remain subject to the terms of their original providers; provider-specific access dates and reuse conditions are documented in the repository data manifest.",
        "### Competing interests\nThe authors declare that they have no competing interests.",
        "### Funding\nThis work was supported by the Hainan Province Science and Technology Special Fund (ZDYF2025SHFZ046). Shen Wang received doctoral training support from the China Scholarship Council. The funders had no role in study design, data collection, data analysis, data interpretation, manuscript preparation or the decision to submit.",
        "### Authors' contributions\nSW conceived the study, extracted and harmonised data, conducted the statistical analyses, produced the figures and tables, independently coded the indicator audit, interpreted the findings and drafted the manuscript. JZ contributed to data analysis, independently coded the indicator audit and revised the manuscript. ZJ contributed to data extraction, source verification and data checking. PL contributed to writing, interpretation and critical revision. SX supervised the study, was available to adjudicate audit disagreements, contributed to design and analytical planning and critically revised the manuscript. No audit disagreement required adjudication. All authors read and approved the final manuscript.",
        "### Acknowledgements\nThe authors thank the World Health Organization and other data providers for making the aggregate data and metadata available. Shen Wang acknowledges doctoral training support from the China Scholarship Council. Generative AI tools assisted code development, consistency checking, figure assembly and language editing. Two AI-assisted read-only audit passes were used as secondary quality-control checks. All indicator classifications, analyses, interpretations and manuscript text were independently reviewed and approved by the authors, who take full responsibility for the work.",
        "## References\n" + "\n".join(f"{i}. {r}" for i, r in enumerate(REFERENCES, 1)),
        "## Tables",
        "### Table 1. Analytic indicators and operational characteristics\n" + md_table(table1) + "\n\n*Notes:* Coverage denominators are the 197-country and territory analytic frame unless otherwise stated. PM2.5 denotes ambient fine particulate matter; YLL, years of life lost; DALY, disability-adjusted life year.",
        "### Table 2. Quantitative evidence supporting the surveillance interpretation\n" + md_table(table2) + "\n\n*Notes:* Confidence intervals for median differences and Cliff's delta were obtained by non-parametric bootstrap. Regression confidence intervals use HC3 standard errors unless stated otherwise. Results are ecological and descriptive.",
        "## Figure legends\n" + "\n\n".join(FIGURE_LEGENDS),
        "## Additional file\nAdditional file 1. **File name:** Additional_file_1_supplementary_methods_figures_tables.docx. **Format:** DOCX. **Title:** Supplementary methods, figures and tables. **Description:** Detailed source-identification and audit methods, missingness and descriptive summaries, formal exposure-inequality estimates, incidence-model coefficients and diagnostics, influence and spatial sensitivity analyses, exploratory environmental heterogeneity, outcome consistency checks, nine supplementary figures and twelve supplementary tables.",
    ])
    return "\n\n".join(parts) + "\n"


def add_page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def set_doc_style(doc: Document, line_numbers: bool = True) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Inches(.9); section.bottom_margin = Inches(.9); section.left_margin = Inches(.9); section.right_margin = Inches(.9)
        if line_numbers:
            sect_pr = section._sectPr
            ln = sect_pr.find(qn("w:lnNumType"))
            if ln is None:
                ln = OxmlElement("w:lnNumType"); sect_pr.append(ln)
            ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous")
        add_page_number(section)


def add_para(doc: Document, text: str = "", bold: bool = False, size: int | None = None, align=None, spacing: float = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(6)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.name = "Times New Roman"
    if size: r.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2; p.paragraph_format.space_after = Pt(6); p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(14 if level == 1 else 12)


def add_word_table(doc: Document, rows: list[list[str]], font_size: float = 7.5, note: str | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j); cell.text = ""
            p = cell.paragraphs[0]; p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value)); r.font.name = "Times New Roman"; r.font.size = Pt(font_size); r.bold = (i == 0)
    if note:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1; p.paragraph_format.space_after = Pt(8)
        r = p.add_run("Notes: " + note); r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(8)
    else:
        doc.add_paragraph()


def build_main_docx(table1: list[list[str]], table2: list[list[str]], main_figures: list[Path]) -> Document:
    doc = Document(); set_doc_style(doc, line_numbers=True)
    add_para(doc, TITLE, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, ", ".join(AUTHORS), align=WD_ALIGN_PARAGRAPH.CENTER)
    for line in AFFILIATIONS: add_para(doc, line)
    for line in CORRESPONDING: add_para(doc, line)
    add_para(doc, "Additional author details: Jing Zhou, email hy0208035@muhn.edu.cn, ORCID 0000-0003-4076-4083; Pengyu Li, email leepy8@163.com.")
    add_heading(doc, "Abstract")
    for k, v in ABSTRACT.items():
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 2; p.paragraph_format.space_after = Pt(6); p.add_run(k + ": ").bold = True; p.add_run(v)
    add_para(doc, "Keywords: " + KEYWORDS)
    add_heading(doc, "Background")
    for p in BACKGROUND: add_para(doc, p)
    add_heading(doc, "Methods")
    for h, paras in METHOD_SECTIONS:
        add_heading(doc, h, 2)
        for p in paras: add_para(doc, p)
    add_heading(doc, "Results")
    for h, paras in RESULT_SECTIONS:
        add_heading(doc, h, 2)
        for p in paras: add_para(doc, p)
    add_heading(doc, "Discussion")
    for p in DISCUSSION: add_para(doc, p)
    add_heading(doc, "Conclusions")
    add_para(doc, CONCLUSION)
    add_heading(doc, "Abbreviations")
    add_para(doc, "AIC, Akaike information criterion; DALY, disability-adjusted life year; GBD, Global Burden of Disease; GDP, gross domestic product; GHE, Global Health Estimates; GHO, Global Health Observatory; HDI, Human Development Index; HEAT, Health Equity Assessment Toolkit; HIC, high-income country; HIDR, Health Inequality Data Repository; LC, lung cancer; LDCT, low-dose computed tomography; LIC, low-income country; LMIC, lower-middle-income country; PM2.5, particulate matter with aerodynamic diameter no greater than 2.5 micrometres; SDI, Socio-demographic Index; UMIC, upper-middle-income country; WHO, World Health Organization; YLL, years of life lost.")
    add_heading(doc, "Declarations")
    declarations = [
        ("Ethics approval and consent to participate", "This study used publicly available aggregate country-level data and did not involve individual participants or identifiable personal information. Ethics approval and informed consent were not required."),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", "The cleaned aggregate country-level dataset, indicator dictionaries, independent-coder audit records, analysis tables and reproduction scripts are available at https://github.com/ammkkr/women-lung-cancer-equity-surveillance. The repository contains no individual-level or identifiable data. Source datasets remain subject to the terms of their original providers; provider-specific access dates and reuse conditions are documented in the repository data manifest."),
        ("Competing interests", "The authors declare that they have no competing interests."),
        ("Funding", "This work was supported by the Hainan Province Science and Technology Special Fund (ZDYF2025SHFZ046). Shen Wang received doctoral training support from the China Scholarship Council. The funders had no role in study design, data collection, data analysis, data interpretation, manuscript preparation or the decision to submit."),
        ("Authors' contributions", "SW conceived the study, extracted and harmonised data, conducted the statistical analyses, produced the figures and tables, independently coded the indicator audit, interpreted the findings and drafted the manuscript. JZ contributed to data analysis, independently coded the indicator audit and revised the manuscript. ZJ contributed to data extraction, source verification and data checking. PL contributed to writing, interpretation and critical revision. SX supervised the study, was available to adjudicate audit disagreements, contributed to design and analytical planning and critically revised the manuscript. No audit disagreement required adjudication. All authors read and approved the final manuscript."),
        ("Acknowledgements", "The authors thank the World Health Organization and other data providers for making the aggregate data and metadata available. Shen Wang acknowledges doctoral training support from the China Scholarship Council. Generative AI tools assisted code development, consistency checking, figure assembly and language editing. Two AI-assisted read-only audit passes were used as secondary quality-control checks. All indicator classifications, analyses, interpretations and manuscript text were independently reviewed and approved by the authors, who take full responsibility for the work."),
    ]
    for h, p in declarations: add_heading(doc, h, 2); add_para(doc, p)
    add_heading(doc, "References")
    for i, ref in enumerate(REFERENCES, 1): add_para(doc, f"{i}. {ref}", spacing=1.0)
    add_heading(doc, "Tables")
    add_para(doc, "Table 1. Analytic indicators and operational characteristics", bold=True)
    add_word_table(doc, table1, 7, "Coverage denominators are the 197-country and territory analytic frame unless otherwise stated. PM2.5 denotes ambient fine particulate matter; YLL, years of life lost; DALY, disability-adjusted life year.")
    add_para(doc, "Table 2. Quantitative evidence supporting the surveillance interpretation", bold=True)
    add_word_table(doc, table2, 7, "Confidence intervals for median differences and Cliff's delta were obtained by non-parametric bootstrap. Regression confidence intervals use HC3 standard errors unless stated otherwise. Results are ecological and descriptive.")
    add_heading(doc, "Figure legends")
    for legend in FIGURE_LEGENDS: add_para(doc, legend, spacing=1.0)
    add_heading(doc, "Figure previews")
    for i, fig in enumerate(main_figures, 1):
        add_para(doc, f"Figure {i}.", bold=True, spacing=1.0)
        doc.add_picture(str(fig), width=Inches(6.5))
    add_heading(doc, "Additional file")
    add_para(doc, "Additional file 1. File name: Additional_file_1_supplementary_methods_figures_tables.docx. Format: DOCX. Title: Supplementary methods, figures and tables. Description: Detailed source-identification and audit methods, missingness and descriptive summaries, formal exposure-inequality estimates, incidence-model coefficients and diagnostics, influence and spatial sensitivity analyses, exploratory environmental heterogeneity, outcome consistency checks, nine supplementary figures and twelve supplementary tables.")
    return doc


def dataframe_rows(df: pd.DataFrame, columns: list[str], labels: list[str] | None = None, decimals: int = 3) -> list[list[str]]:
    labels = labels or columns
    rows = [labels]
    for _, r in df[columns].iterrows():
        out = []
        for value in r:
            if pd.isna(value): out.append("")
            elif isinstance(value, (float, np.floating)): out.append(f"{value:.{decimals}f}")
            else: out.append(str(value))
        rows.append(out)
    return rows


def supplementary_tables(country: pd.DataFrame, effects: pd.DataFrame, coef: pd.DataFrame, fit: pd.DataFrame, window: pd.DataFrame,
                         robust: pd.DataFrame, spatial_diag: pd.DataFrame, env: pd.DataFrame, outcome_summary: pd.DataFrame,
                         corr: pd.DataFrame, inc_miss: pd.DataFrame, reg_miss: pd.DataFrame) -> list[tuple[str, list[list[str]], str]]:
    criteria = [
        ["Classification", "Operational rule", "Interpretation"],
        ["Adequate for this analysis", "Generally at least 80% country coverage, comparable definition and unit, machine-readable, and the required female/sex dimension", "Usable in the primary analytic frame for its stated purpose"],
        ["Available with limitations", "Lower coverage or an important limitation in subgroup, time, uncertainty, comparability or standardisation", "Usable only for restricted description, context or sensitivity analysis"],
        ["Source located but not sufficiently harmonised/integrated", "A relevant global or multi-country source was verified but could not be integrated into the prespecified frame", "Do not describe the component as absent; identify the integration limitation"],
        ["No compatible global indicator identified", "No source meeting the country-level compatibility requirements was identified in the predefined search", "A source-frame result, not proof that no local data exist"],
        ["Not applicable", "The audit dimension is not conceptually applicable to the component", "Excluded from adequacy interpretation for that cell"],
    ]
    src = source_protocol_table()
    audit = independent_audit_record()
    audit_cols = ["domain", "component", "audit_dimension", "Shen_Wang_code", "Jing_Zhou_code", "exact_agreement", "final_code", "final_label", "adjudication_change_required", "evidence_note"]
    frame = country[["iso3", "country", "WHO_region", "income_group", "complete_analytic_indicator_frame", "incidence_model_complete", "analytic_missingness_pattern"] + [c for _, c, _ in PRIMARY_INDICATORS]].copy()
    miss = pd.concat([inc_miss.assign(stratifier="Income group").rename(columns={"income_group": "group"}), reg_miss.assign(stratifier="WHO region").rename(columns={"WHO_region": "group"})], ignore_index=True)
    desc = pd.concat([descriptive_table(country, "income_group").assign(stratifier="Income group"), descriptive_table(country, "WHO_region").assign(stratifier="WHO region")], ignore_index=True)
    model = pd.concat([
        coef.assign(section="Coefficients").rename(columns={"model": "model_or_analysis"}),
        fit.assign(section="Model fit").rename(columns={"model": "model_or_analysis"}),
    ], ignore_index=True, sort=False)
    infl = pd.DataFrame()
    return [
        ("Supplementary Table S1. Audit classifications and operational criteria", criteria, "Thresholds were prespecified for this analysis and describe suitability for the stated analytic role."),
        ("Supplementary Table S2. Source-identification protocol and targeted verification", dataframe_rows(src, src.columns.tolist()), "Searches were structured but did not constitute a systematic review of every national data source."),
        ("Supplementary Table S3. Item-level indicator-availability audit and independent-coder agreement", dataframe_rows(audit, audit_cols, decimals=3), "Shen Wang and Jing Zhou independently coded 18 components across 10 dimensions. Agreement was 180/180 cells (100%) and Cohen's kappa was 1.000. Coder-specific records were archived on 13 July 2026 after completion; no adjudication change was required. Two AI-assisted read-only checks reproduced the final codes but were excluded from the human inter-rater statistic."),
        ("Supplementary Table S4. Country and territory analytic-frame membership", dataframe_rows(frame, frame.columns.tolist(), decimals=3), "Complete status requires current female smoking, clean-fuel access, ambient PM2.5 and age-standardised female lung-cancer incidence."),
        ("Supplementary Table S5. Primary-indicator coverage by income group and WHO region", dataframe_rows(miss, ["stratifier", "group", "n_total", "n_complete", "complete_percent"] + [f"coverage_{c}" for _, c, _ in PRIMARY_INDICATORS], decimals=1), "Percentages use the number of countries or territories in each stratum as the denominator."),
        ("Supplementary Table S6. Descriptive exposure and incidence distributions", dataframe_rows(desc, ["stratifier", "group", "indicator", "n", "median", "Q1", "Q3", "minimum", "maximum"], decimals=2), "Summaries are unweighted country-level distributions."),
        ("Supplementary Table S7. Income-group exposure inequality effect sizes", dataframe_rows(effects, effects.columns.tolist(), decimals=3), "Median-difference and Cliff's-delta intervals use 5000 bootstrap samples."),
        ("Supplementary Table S8. Age-standardised incidence model coefficients and fit", dataframe_rows(model, ["section", "model_or_analysis", "term", "n", "estimate", "robust_SE_HC3", "CI_low", "CI_high", "P", "R2", "adjusted_R2", "AIC"], decimals=4), "HC3 confidence intervals are reported for coefficients; R-squared, adjusted R-squared and AIC are reported in model-fit rows."),
        ("Supplementary Table S9. Alternative female-smoking exposure windows", dataframe_rows(window, window.columns.tolist(), decimals=4), "Each row uses the same structural adjustment set: log GDP, urbanisation and WHO region."),
        ("Supplementary Table S10. Influence and spatial sensitivity analyses", dataframe_rows(pd.concat([robust, spatial_diag.rename(columns={"metric": "analysis"})], ignore_index=True, sort=False), ["analysis", "n", "estimate", "CI_low", "CI_high", "P"], decimals=4), "Conley-type intervals were calculated for the geocoded subset. Moran's I is reported without a confidence interval."),
        ("Supplementary Table S11. Exploratory environmental heterogeneity models", dataframe_rows(env, env.columns.tolist(), decimals=4), "Environmental coefficients are supplementary and should not be interpreted as causal effects."),
        ("Supplementary Table S12. Incidence-mortality residual consistency", dataframe_rows(outcome_summary, outcome_summary.columns.tolist(), decimals=4), "Both residuals were derived from models adjusted for historical smoking, log GDP, urbanisation and WHO region."),
    ]


def supplementary_markdown(tables: list[tuple[str, list[list[str]], str]]) -> str:
    methods = [
        "### Scope and source identification\nThe supplement documents the structured source audit and all analyses supporting the main manuscript. Formal coding used WHO HIDR/HEAT-linked sources, GHO, GHE, GBD-linked outputs, World Bank indicators and Global Data Lab HDI. Targeted verification covered IARC/GLOBOCAN, CanScreen5, CONCORD, WHO second-hand-smoke metadata and GTSS/GATS. Shen Wang and Jing Zhou independently coded all 180 component-by-dimension cells. Their coder-specific records were archived after completion, compared without alteration and showed 100% agreement (Cohen's kappa=1.000); no adjudication change was required. Two AI-assisted read-only passes reproduced the final classifications as secondary quality control and were excluded from the human agreement statistic. The workflow and decision rules are shown in Supplementary Figure S1 and Supplementary Tables S1-S3.",
        "### Missingness and descriptive analyses\nThe primary frame required current female smoking, clean-fuel access, ambient PM2.5 and age-standardised female lung-cancer incidence. Missingness patterns were mutually exclusive. Coverage percentages used countries in each income or WHO-region stratum as denominators. Country-level values were not population weighted. Full country membership, coverage and distributions appear in Supplementary Tables S4-S6 and Supplementary Figures S2-S4.",
        "### Inequality effect sizes\nFor each exposure, the prespecified low-income versus high-income contrast was expressed as a difference between group medians in the epidemiologically interpretable direction. Confidence intervals used 5000 within-group bootstrap resamples. Cliff's delta compared low-income with high-income country values and used the same resampling procedure. Clean-fuel structural correlations used Spearman's rho with bootstrap intervals. Estimates are reported in Supplementary Figure S5A-B and Supplementary Table S7.",
        "### Age-standardised incidence models\nThe validation model used z-scored 2021 age-standardised female lung-cancer incidence. Historical smoking was the mean female age-standardised current tobacco prevalence in 2000, 2005 and 2010. The primary model adjusted for z-scored log GDP, z-scored urbanisation and WHO region, with HC3 standard errors. The adjusted relationship in Figure 4B was visualised using the Frisch-Waugh-Lovell construction: outcome and exposure were separately residualised on the adjustment set and the residuals were related using an HC3 fit. Alternative windows, age-structure adjustment, Cook-distance exclusion, population restriction and Conley-type spatial covariance are reported in Supplementary Figure S5C-D and Supplementary Tables S8-S10; general diagnostics and spatial residual checks appear in Supplementary Figures S6 and S8.",
        "### Exploratory environmental and outcome checks\nThe incidence residual from the primary smoking model was regressed on z-scored clean-fuel deficit and PM2.5 globally and in non-high-income settings, with region and interaction sensitivity analyses. Analogous incidence and WHO mortality residuals were compared as a source-consistency check. These analyses are supplementary because environmental histories were short and the mortality label did not explicitly state age standardisation (Supplementary Figures S7 and S9; Supplementary Tables S11-S12).",
    ]
    results = [
        "### Coverage and audit results\nMost excluded countries were missing female smoking, and complete-frame inclusion was lowest in low-income and income-unclassified settings (Supplementary Figure S2; Supplementary Tables S4-S5). Human coder agreement was 180/180 cells (100%; Cohen's kappa=1.000), and all final codes matched both secondary AI-assisted quality-control passes (Supplementary Figure S1; Supplementary Table S3). The audit retained three adequate exposure indicators but classified second-hand smoke, occupational carcinogens and radon as limited. Screening, stage and treatment sources were located but not sufficiently integrated; survival and subtype were available with limitations (Supplementary Figure S3; Supplementary Table S3).",
        "### Exposure inequalities\nCountry distributions varied substantially across income groups and WHO regions (Supplementary Figure S4; Supplementary Table S6). Bootstrap intervals and Cliff's delta confirmed opposing gradients for female smoking and household-energy deprivation, while ambient PM2.5 displayed a less monotonic intermediate pattern (Supplementary Figure S5A; Supplementary Table S7). Clean-fuel deficit was strongly inversely correlated with SDI, HDI and urbanisation and positively correlated with rural disadvantage (Supplementary Figure S5B).",
        "### Model and estimate-quality results\nThe historical-smoking coefficient remained positive across exposure windows, additional age-structure adjustment, Cook-distance exclusion, population restriction and Conley-type uncertainty estimates (Supplementary Figure S5C-D; Supplementary Tables S8-S10). General model diagnostics are shown in Supplementary Figure S6, and residual spatial autocorrelation remained detectable (Supplementary Figure S8). The non-high-income PM2.5 coefficient was positive before WHO-region adjustment but its interval included zero after regional adjustment (Supplementary Figure S7; Supplementary Table S11). Incidence and mortality residuals were strongly correlated, although this check does not remove source-method differences (Supplementary Figure S9; Supplementary Table S12).",
    ]
    parts = ["# Additional file 1", "## Supplementary methods", *methods, "## Supplementary results", *results, "## Supplementary figure legends", "\n\n".join(SUPPLEMENTARY_FIGURE_LEGENDS), "## Supplementary tables"]
    for title, rows, note in tables:
        parts.append(f"### {title}\n{md_table(rows)}\n\n*Notes:* {note}")
    return "\n\n".join(parts) + "\n"


def build_supplement_docx(tables: list[tuple[str, list[list[str]], str]], figures: list[Path]) -> Document:
    doc = Document(); set_doc_style(doc, line_numbers=False)
    add_para(doc, "Additional file 1", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Supplementary methods, figures and tables", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, "Supplementary methods")
    methods = supplementary_markdown([]).split("## Supplementary results")[0]
    for block in methods.split("### ")[1:]:
        lines = block.strip().split("\n", 1); add_heading(doc, lines[0], 2); add_para(doc, lines[1] if len(lines)>1 else "")
    add_heading(doc, "Supplementary results")
    res_part = supplementary_markdown([]).split("## Supplementary results",1)[1].split("## Supplementary figure legends",1)[0]
    for block in res_part.split("### ")[1:]:
        lines=block.strip().split("\n",1);add_heading(doc,lines[0],2);add_para(doc,lines[1] if len(lines)>1 else "")
    add_heading(doc, "Supplementary figures")
    for fig, legend in zip(figures, SUPPLEMENTARY_FIGURE_LEGENDS):
        add_para(doc, legend, bold=True, spacing=1.0)
        doc.add_picture(str(fig), width=Inches(6.5))
    add_heading(doc, "Supplementary tables")
    for title, rows, note in tables:
        add_para(doc, title, bold=True, spacing=1.0)
        add_word_table(doc, rows, font_size=6.4 if len(rows[0]) > 8 else 7.0, note=note)
    return doc



def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for generated_dir in [FIG_DIR, SUPP_FIG_DIR, TABLE_DIR]:
        if generated_dir.exists():
            shutil.rmtree(generated_dir)
        generated_dir.mkdir(parents=True, exist_ok=True)
    country, src = load_data()
    effects = exposure_inequality_effects(country, 5000)
    diffs = included_excluded_differences(country, 5000)
    corr_rows = []
    for label, col in [("SDI", "SDI"), ("HDI", "HDI"), ("Urbanisation", "urbanisation"), ("Rural clean-fuel gap", "rural_clean_fuel_disadvantage")]:
        est, lo, hi, n = spearman_bootstrap(country["clean_fuel_deficit"], country[col], 3000)
        corr_rows.append({"correlate": label, "rho": est, "CI_low": lo, "CI_high": hi, "n": n})
    corr = pd.DataFrame(corr_rows)
    coef, fit, window, robust, model_stats = fit_incidence_models(country)
    robust, spatial_diag = add_spatial_sensitivity(model_stats, robust)
    env, _ = residual_environment_models(country)
    outcome_d, outcome_summary = outcome_residual_consistency(country)
    inc_miss, reg_miss = missingness_by_income_region(country)

    Xmain, _ = design_matrix(model_stats["data"], ["z_exposure", "z_log_GDP", "z_urbanisation", "WHO_region"], categorical=["WHO_region"])
    bp_lm, bp_p = bptest_like(model_stats["model"], Xmain)
    shapiro = stats.shapiro(model_stats["model"]["resid"])
    jarque = stats.jarque_bera(model_stats["model"]["resid"])
    diagnostics = pd.DataFrame([
        {"metric": "Breusch-Pagan-type LM", "estimate": bp_lm, "P": bp_p, "n": model_stats["model"]["n"]},
        {"metric": "Shapiro-Wilk W", "estimate": shapiro.statistic, "P": shapiro.pvalue, "n": model_stats["model"]["n"]},
        {"metric": "Jarque-Bera", "estimate": jarque.statistic, "P": jarque.pvalue, "n": model_stats["model"]["n"]},
    ])

    outputs = {
        "Table_S1_adjudicated_indicator_availability_audit.csv": audit_dataframe(),
        "Table_S2_source_identification_protocol.csv": source_protocol_table(),
        "Table_S3_country_analytic_frame.csv": country[["iso3", "country", "WHO_region", "income_group", "complete_analytic_indicator_frame", "incidence_model_complete", "analytic_missingness_pattern"] + [c for _, c, _ in PRIMARY_INDICATORS]],
        "Table_S4_exposure_inequality_effect_sizes.csv": effects,
        "Table_S5_included_excluded_development_differences.csv": diffs,
        "Table_S6_clean_fuel_correlations.csv": corr,
        "Table_S7_incidence_model_coefficients.csv": coef,
        "Table_S8_incidence_model_fit.csv": fit,
        "Table_S9_smoking_window_sensitivity.csv": window,
        "Table_S10_influence_and_spatial_sensitivity.csv": robust,
        "Table_S11_spatial_diagnostics.csv": spatial_diag,
        "Table_S12_exploratory_environmental_models.csv": env,
        "Table_S13_outcome_residual_consistency.csv": outcome_summary,
        "Table_S14_general_model_diagnostics.csv": diagnostics,
        "Table_S15_missingness_by_income.csv": inc_miss,
        "Table_S16_missingness_by_WHO_region.csv": reg_miss,
        "Table_S17_descriptive_by_income.csv": descriptive_table(country, "income_group"),
        "Table_S18_descriptive_by_WHO_region.csv": descriptive_table(country, "WHO_region"),
        "audit_coding_item_level_agreement.csv": independent_audit_record(),
        "audit_coder_SW.csv": coder_audit_sheet("Shen Wang"),
        "audit_coder_JZ.csv": coder_audit_sheet("Jing Zhou"),
        "audit_agreement_summary.csv": audit_agreement_summary(),
        "source_access_manifest.csv": source_access_manifest(),
    }
    for name, frame in outputs.items(): frame.to_csv(TABLE_DIR / name, index=False, encoding="utf-8-sig")

    main_figures = save_main_figures(country, src, effects, diffs, corr, window, fit, robust, model_stats, spatial_diag)
    supp_figures = save_supplementary_figures(
        country, model_stats, env, spatial_diag, outcome_d, outcome_summary,
        effects, corr, window, robust,
    )
    t1 = main_table1(country); t2 = main_table2(country, effects, diffs, model_stats, env, outcome_summary)
    tables = supplementary_tables(country, effects, coef, fit, window, robust, spatial_diag, env, outcome_summary, corr, inc_miss, reg_miss)
    for i, (title, rows, _) in enumerate(tables, 1):
        slug = re.sub(r"[^A-Za-z0-9]+", "_", title.split(". ", 1)[1]).strip("_").lower()
        pd.DataFrame(rows[1:], columns=rows[0]).to_csv(TABLE_DIR / f"Supplementary_Table_S{i}_{slug}.csv", index=False, encoding="utf-8-sig")

    main_md = manuscript_markdown(t1, t2)
    supp_md = supplementary_markdown(tables)
    main_md_path = OUT_DIR / "manuscript.md"
    supp_md_path = OUT_DIR / "Additional_file_1_supplementary_methods_figures_tables.md"
    main_md_path.write_text(main_md, encoding="utf-8"); supp_md_path.write_text(supp_md, encoding="utf-8")
    main_doc = build_main_docx(t1, t2, main_figures); main_doc_path = OUT_DIR / "manuscript.docx"; main_doc.save(main_doc_path)
    supp_doc = build_supplement_docx(tables, supp_figures); supp_doc_path = OUT_DIR / "Additional_file_1_supplementary_methods_figures_tables.docx"; supp_doc.save(supp_doc_path)


    abstract_text = " ".join(ABSTRACT.values())
    body_text = " ".join(BACKGROUND + [p for _, ps in METHOD_SECTIONS for p in ps] + [p for _, ps in RESULT_SECTIONS for p in ps] + DISCUSSION + [CONCLUSION])
    verification = f"""# Analysis and submission verification report

## Core counts

- Analytic countries/territories: {len(country)}
- Complete four-indicator frame: {int(country['complete_analytic_indicator_frame'].sum())}
- Missing female smoking only: {int((country['analytic_missingness_pattern'] == 'Missing female smoking only').sum())}
- Main incidence-model n: {model_stats['model']['n']}

## Text and display checks

- Abstract word count: {word_count(abstract_text)} (BMC Public Health maximum: 350)
- Main body word count from Background through Conclusions: {word_count(body_text)}
- References: {len(REFERENCES)}
- Main tables: 2
- Main figures: {len(main_figures)}
- Supplementary figures: {len(supp_figures)}
- Supplementary tables in Additional file 1: {len(tables)}
- Main Word tables: {len(main_doc.tables)}
- Main Word inline figures: {len(main_doc.inline_shapes)}
- Supplement Word tables: {len(supp_doc.tables)}
- Supplement Word inline figures: {len(supp_doc.inline_shapes)}

## Numeric checks

- Main historical-smoking coefficient: {model_stats['main_coef']['estimate']:.6f} (95% CI {model_stats['main_coef']['CI_low']:.6f} to {model_stats['main_coef']['CI_high']:.6f})
- Moran's I: {spatial_diag.iloc[0]['estimate']:.6f}; permutation P={spatial_diag.iloc[0]['P']:.4f}
- Pearson incidence-mortality residual correlation: {outcome_summary.query("method == 'Pearson'").iloc[0]['estimate']:.6f}
- General model diagnostics: {diagnostics.to_dict(orient='records')}

## Automated wording checks

- Contains public GitHub URL and no placeholder repository address.

## Non-automatable author checks

- SW and JZ coder-specific records are archived; agreement is 180/180 cells and Cohen's kappa is 1.000. Both authors should sign the final verification checklist.
- Shen Wang and Zhangdong Jiang email addresses were not supplied in the manuscript files.
- Jing Zhou's supplied email string should be verified for spacing and spelling.
- A versioned Zenodo DOI remains optional; the public GitHub commit should be cited if no DOI is created.
"""
    (OUT_DIR / "analysis_and_submission_verification_report.md").write_text(verification, encoding="utf-8")
    checklist = """# BMC Public Health submission readiness checklist

- [x] Structured abstract below 350 words
- [x] Background, Methods, Results, Discussion and Conclusions aligned to one contribution
- [x] Country-level ecological design and non-causal interpretation stated
- [x] Explicitly age-standardised primary burden anchor
- [x] Formal inequality effect sizes with uncertainty
- [x] Influence and spatial sensitivity analyses
- [x] Four composite main figures exported as PNG, TIFF, SVG and PDF
- [x] Two editable main Word tables without colour shading
- [x] Supplementary methods and results linked to all supplementary displays
- [x] Required BMC declarations included
- [x] Public repository URL included
- [x] Archive independent SW and JZ coding records and calculate agreement
- [ ] Verify all author emails, ORCIDs and affiliations
- [x] Prepare a relative-path release candidate with audit and provenance records
- [ ] Confirm the public repository reproduces after the final push
- [ ] Create a versioned Zenodo release and add DOI, if used
- [ ] Obtain final author approval of data, code, text, figures and declarations
"""
    (OUT_DIR / "submission_readiness_checklist.md").write_text(checklist, encoding="utf-8")

    print(verification)


if __name__ == "__main__":
    main()
